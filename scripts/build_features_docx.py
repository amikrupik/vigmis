# -*- coding: utf-8 -*-
"""
Build VIGMIS_FEATURES_HE.docx from VIGMIS_FEATURES_HE.md
Run: python scripts/build_features_docx.py
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
MD_PATH  = ROOT / "docs" / "VIGMIS_FEATURES_HE.md"
OUT_PATH = ROOT / "docs" / "VIGMIS_FEATURES_HE.docx"

# ── colour palette ──────────────────────────────────────────────────────────
C_TITLE    = RGBColor(0x1a, 0x1a, 0x2e)   # near-black navy
C_H1       = RGBColor(0x16, 0x21, 0x3e)   # deep navy
C_H2       = RGBColor(0x0f, 0x3d, 0x6e)   # medium navy
C_H3       = RGBColor(0x1a, 0x73, 0xe8)   # google-blue
C_H4       = RGBColor(0x34, 0xa8, 0x53)   # google-green
C_CODE_BG  = RGBColor(0xf4, 0xf4, 0xf4)
C_TABLE_H  = RGBColor(0x1a, 0x73, 0xe8)
C_TABLE_AH = RGBColor(0xFF, 0xFF, 0xFF)
C_RULE     = RGBColor(0xcc, 0xcc, 0xcc)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def add_heading(doc, text, level):
    colours = {1: C_H1, 2: C_H2, 3: C_H3, 4: C_H4}
    sizes   = {1: 22,   2: 17,   3: 14,   4: 12}
    bolds   = {1: True, 2: True, 3: True, 4: False}

    # strip leading ## markers if present
    text = re.sub(r'^#+\s*', '', text).strip()
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = colours.get(level, C_H2)
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = bolds.get(level, False)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def inline_format(run_text: str):
    """Return list of (text, bold, italic, code) tuples."""
    parts = []
    # split on **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    last = 0
    for m in pattern.finditer(run_text):
        if m.start() > last:
            parts.append((run_text[last:m.start()], False, False, False))
        tok = m.group()
        if tok.startswith('**'):
            parts.append((tok[2:-2], True, False, False))
        elif tok.startswith('*'):
            parts.append((tok[1:-1], False, True, False))
        else:  # backtick
            parts.append((tok[1:-1], False, False, True))
        last = m.end()
    if last < len(run_text):
        parts.append((run_text[last:], False, False, False))
    return parts

def add_para(doc, text: str, indent=0, bullet=False, bold=False, italic=False):
    # strip leading list markers
    clean = re.sub(r'^[-*]\s+', '', text).strip()
    clean = re.sub(r'^\d+\.\s+', '', clean)

    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()

    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.25)

    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

    parts = inline_format(clean)
    for txt, b, i, code in parts:
        r = p.add_run(txt)
        r.font.bold   = b or bold
        r.font.italic = i or italic
        r.font.size   = Pt(10.5)
        if code:
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xd6, 0x33, 0x6c)
    return p

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line)
        r.font.name  = 'Courier New'
        r.font.size  = Pt(9)
        r.font.color.rgb = RGBColor(0x24, 0x29, 0x2e)
        # light grey shade on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F4F4F4')
        pPr.append(shd)

def add_table(doc, header_row, data_rows):
    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header
    hdr = table.rows[0]
    for i, cell_text in enumerate(header_row):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1a73e8')
        p = cell.paragraphs[0]
        r = p.add_run(cell_text.strip())
        r.font.bold  = True
        r.font.color.rgb = C_TABLE_AH
        r.font.size  = Pt(10)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # data rows
    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'F0F5FF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            parts = inline_format(cell_text.strip())
            for txt, b, i, code in parts:
                r = p.add_run(txt)
                r.font.size = Pt(9.5)
                r.font.bold = b
                r.font.italic = i
                if code:
                    r.font.name = 'Courier New'
                    r.font.color.rgb = RGBColor(0xd6, 0x33, 0x6c)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()  # spacer


def parse_table_line(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def build_doc():
    doc = Document()

    # ── page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── default paragraph font ──────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    # ── title page ──────────────────────────────────────────────────────────
    tp = doc.add_heading('Vigmis — רשימת פיצ׳רים מלאה', level=0)
    tp.runs[0].font.color.rgb = C_TITLE
    tp.runs[0].font.size = Pt(28)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('מסמך מלא של כל היכולות, הפיצ׳רים, האינטגרציות והתוכנות המשתתפות')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = C_H2
    date_p = doc.add_paragraph('תאריך עדכון: 2026‑05‑30')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # ── parse markdown ──────────────────────────────────────────────────────
    text = MD_PATH.read_text(encoding='utf-8')
    lines = text.splitlines()

    in_code = False
    code_buf = []
    in_table = False
    table_header = []
    table_rows = []
    skip_sep = False

    i = 0
    while i < len(lines):
        line = lines[i]
        raw  = line

        # ── code fence ──────────────────────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                add_code_block(doc, code_buf)
                in_code = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── markdown table ───────────────────────────────────────────────────
        if line.strip().startswith('|'):
            if not in_table:
                in_table    = True
                table_header = parse_table_line(line)
                skip_sep     = True
            elif skip_sep and re.match(r'^\|[-| :]+\|', line.strip()):
                skip_sep = False  # separator row
            else:
                table_rows.append(parse_table_line(line))
            i += 1
            continue
        else:
            if in_table:
                add_table(doc, table_header, table_rows)
                in_table = False
                table_header = []
                table_rows   = []
                skip_sep     = False

        stripped = line.strip()

        # ── horizontal rule ─────────────────────────────────────────────────
        if stripped in ('---', '***', '___') or re.match(r'^-{3,}$', stripped):
            add_hrule(doc)
            i += 1
            continue

        # ── title (# level 1) ───────────────────────────────────────────────
        if stripped.startswith('#### '):
            add_heading(doc, stripped[5:], 4)
        elif stripped.startswith('### '):
            add_heading(doc, stripped[4:], 3)
        elif stripped.startswith('## '):
            add_heading(doc, stripped[3:], 2)
        elif stripped.startswith('# '):
            add_heading(doc, stripped[2:], 1)

        # ── bullet / numbered list ──────────────────────────────────────────
        elif re.match(r'^[-*]\s', stripped) or re.match(r'^\d+\.\s', stripped):
            add_para(doc, stripped, bullet=True)

        # ── blockquote ──────────────────────────────────────────────────────
        elif stripped.startswith('>'):
            p = add_para(doc, stripped[1:].strip(), indent=1, italic=True)
            p.paragraph_format.left_indent = Inches(0.3)

        # ── italic star line (e.g. *סוף המסמך.*) ──────────────────────────
        elif stripped.startswith('*') and stripped.endswith('*') and stripped.count('*') == 2:
            add_para(doc, stripped[1:-1], italic=True)

        # ── empty line ──────────────────────────────────────────────────────
        elif stripped == '':
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ── regular paragraph ───────────────────────────────────────────────
        else:
            add_para(doc, stripped)

        i += 1

    # flush any trailing table
    if in_table:
        add_table(doc, table_header, table_rows)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")

if __name__ == '__main__':
    build_doc()
