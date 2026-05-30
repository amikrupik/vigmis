# -*- coding: utf-8 -*-
"""
Build QA-Ana-HaAretzHaTova.docx — Hebrew QA checklist for real user testing
Run: python scripts/build_ana_qa_docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = Path(__file__).parent.parent / "qa" / "QA-Ana-HaAretzHaTova.docx"

C_NAVY   = RGBColor(0x1a, 0x21, 0x3e)
C_BLUE   = RGBColor(0x1a, 0x73, 0xe8)
C_GREEN  = RGBColor(0x0d, 0x90, 0x4f)
C_RED    = RGBColor(0xd9, 0x30, 0x25)
C_ORANGE = RGBColor(0xe3, 0x7a, 0x00)
C_GREY   = RGBColor(0x55, 0x55, 0x55)
C_LGREY  = RGBColor(0xf8, 0xf9, 0xfa)

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_break(doc):
    doc.add_page_break()

def heading(doc, text, level=1, color=None, rtl=True):
    sizes = {0: 26, 1: 18, 2: 14, 3: 12}
    colors = {0: C_NAVY, 1: C_NAVY, 2: C_BLUE, 3: C_GREEN}
    p = doc.add_heading(text, level=min(level, 3))
    if p.runs:
        r = p.runs[0]
        r.font.color.rgb = color or colors.get(level, C_NAVY)
        r.font.size = Pt(sizes.get(level, 12))
        r.font.bold = level <= 2
    p.paragraph_format.space_before = Pt(16 if level <= 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    if rtl:
        set_rtl(p)
    return p

def para(doc, text, bold=False, italic=False, color=None, size=11, indent=0, rtl=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    if rtl:
        set_rtl(p)
    return p

def checkbox(doc, text, bold=False, indent=0.2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('☐  ')
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.bold = bold
    set_rtl(p)
    return p

def step_header(doc, number, title, duration=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'שלב {number}: {title}')
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = C_BLUE
    if duration:
        r2 = p.add_run(f'  ({duration})')
        r2.font.size = Pt(10)
        r2.font.color.rgb = C_GREY
        r2.font.italic = True
    set_rtl(p)

def action_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('▶  ')
    r.font.color.rgb = C_BLUE
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.bold = True
    set_rtl(p)

def expect_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('✓  ')
    r.font.color.rgb = C_GREEN
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    set_rtl(p)

def report_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('⚠  ')
    r.font.color.rgb = C_RED
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = C_RED
    set_rtl(p)

def note_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF3CD')
    pPr.append(shd)
    r = p.add_run('💡  ' + text)
    r.font.size = Pt(10)
    r.font.italic = True
    set_rtl(p)

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pb.append(bottom); pPr.append(pb)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def phase_title(doc, number, title, duration):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1a213e')
    pPr.append(shd)
    r = p.add_run(f'  שלב {number} — {title}  ({duration})')
    r.font.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_rtl(p)

def bug_report_template(doc):
    heading(doc, 'טופס דיווח תקלה', level=2, color=C_RED)
    note_box(doc, 'כל פעם שמשהו לא עובד כמצופה — ממלאים את הטופס הבא ושולחים')
    doc.add_paragraph()
    fields = [
        ('שם הבדיקה שנכשלה', ''),
        ('מה עשיתי', ''),
        ('מה ציפיתי לראות', ''),
        ('מה קיבלתי בפועל', ''),
        ('צילום מסך', '□ צורף   □ לא צורף'),
        ('חומרה', '□ קריטי   □ גבוה   □ בינוני   □ נמוך'),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        set_cell_bg(row.cells[0], 'E8F0FE')
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True; r0.font.size = Pt(10)
        set_rtl(p0)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.size = Pt(10)
        set_rtl(p1)
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.5)
    doc.add_paragraph()

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
        section.page_width = Cm(21); section.page_height = Cm(29.7)

    style = doc.styles['Normal']
    style.font.name = 'Arial'; style.font.size = Pt(11)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    tp = doc.add_heading('ספר בדיקות QA — אנה / הארץ הטובה', level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if tp.runs: tp.runs[0].font.color.rgb = C_NAVY; tp.runs[0].font.size = Pt(24)
    set_rtl(tp)

    sub = doc.add_paragraph('מדריך בדיקות מפורט לבדיקת פלטפורמת Vigmis')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs: sub.runs[0].font.size = Pt(13); sub.runs[0].font.color.rgb = C_BLUE
    set_rtl(sub)

    doc.add_paragraph()
    note_box(doc, 'מסמך זה מיועד לבדיקה ידנית מלאה של המערכת. יש לעבור על כל שלב לפי הסדר ולסמן ✓ כל פריט שעובד, או למלא טופס תקלה כשמשהו לא תקין.')
    doc.add_paragraph()

    # Persona table
    heading(doc, 'פרופיל הבודקת', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    rows_data = [
        ('שם', 'אנה'),
        ('עסק', 'הארץ הטובה — חנות אורגנית אונליין'),
        ('אתר', 'haartezhatova.co.il'),
        ('פלטפורמות פרסום', 'Google Ads פעיל + דף Facebook + Instagram'),
        ('תקציב', '₪8,000 בחודש (~$2,200)'),
        ('מטרה', 'יותר רכישות — ROAS > 4'),
        ('תוכנית', 'Free (שדרוג ל-Pro במהלך הבדיקה)'),
    ]
    for i, (k, v) in enumerate(rows_data):
        set_cell_bg(table.rows[i].cells[0], '1a73e8')
        p0 = table.rows[i].cells[0].paragraphs[0]
        r0 = p0.add_run(k); r0.font.bold = True; r0.font.color.rgb = RGBColor(255,255,255); r0.font.size = Pt(10)
        set_rtl(p0)
        p1 = table.rows[i].cells[1].paragraphs[0]
        r1 = p1.add_run(v); r1.font.size = Pt(10)
        set_rtl(p1)
    doc.add_paragraph()

    # Duration
    heading(doc, 'זמן בדיקה משוער: כ-3.5 שעות', level=3, color=C_GREY)
    divider(doc)
    add_page_break(doc)

    # ── PHASE 1 ─────────────────────────────────────────────────────────────
    phase_title(doc, 1, 'כניסה ראשונה + הרשמה', '30 דקות')

    step_header(doc, '1.1', 'כניסה לדף הנחיתה')
    action_line(doc, 'היכנסי לאתר vigmis.com')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'הדף נטען ומוצג תוך 3 שניות')
    expect_line(doc, 'יש כפתור "התחילי בחינם" ברור')
    expect_line(doc, 'מופיע מחיר / עמוד תמחור')
    expect_line(doc, 'יש שאלות ותשובות (FAQ)')
    para(doc, 'מה לדווח אם לא עובד:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'הדף לא נטען כלל')
    report_line(doc, 'כפתור "התחילי בחינם" מוביל לשגיאה')
    report_line(doc, 'אין מחירים / מחירים שגויים')

    divider(doc)
    step_header(doc, '1.2', 'הרשמה')
    action_line(doc, 'לחצי על "התחילי בחינם" ← טופס הרשמה')
    action_line(doc, 'הירשמי עם כתובת מייל + סיסמה')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'טופס הרשמה מופיע עם שדה מייל + סיסמה')
    expect_line(doc, 'אפשרות להירשם עם Google')
    expect_line(doc, 'תוך 2 דקות — מייל אימות מגיע לתיבת הדואר')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'הטופס לא נשלח')
    report_line(doc, 'לא מגיע מייל אימות תוך 2 דקות')
    report_line(doc, 'אחרי אימות — מועברת לדף שגוי')

    divider(doc)
    step_header(doc, '1.3', 'אחרי ההרשמה — מעבר ל-Onboarding')
    action_line(doc, 'אמתי את כתובת המייל ← חזרי לאפליקציה')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'מועברת אוטומטית לעמוד Onboarding')
    expect_line(doc, 'מוצגת הודעת ברוכה הבאה')
    expect_line(doc, 'יש ממשק שיחה / צ׳אט לקבלת פרטים')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'מועברת ישירות ל-Dashboard ומדלגת על Onboarding')
    report_line(doc, 'עמוד ריק או טעינה אינסופית')

    add_page_break(doc)

    # ── PHASE 2 ─────────────────────────────────────────────────────────────
    phase_title(doc, 2, 'שאלון קליטה (Onboarding)', '30 דקות')
    note_box(doc, 'יש 10 נושאים שחייבים להיות מכוסים. המערכת מנהלת שיחה — ענו בטבעיות.')

    topics = [
        ('2.1', 'סוג העסק',
         '"אני מוכרת מוצרי מזון אורגני אונליין, בעיקר דרך האתר שלי"',
         'המערכת מאשרת: "הבנתי — חנות e-commerce"', 'המערכת לא מבינה את סוג העסק'),
        ('2.2', 'כתובת האתר',
         '"האתר שלי הוא haartezhatova.co.il"',
         'המערכת אומרת שתנתח את האתר', 'המערכת לא שומרת את הכתובת'),
        ('2.3', 'מטרת הפרסום + מרווח רווח',
         '"אני רוצה יותר רכישות, המרווח שלי כ-35%"',
         'המערכת מאשרת: מטרה = רכישות, מרווח = 35%', 'מטרה נשמרת כ"לידים" או "תנועה" בטעות'),
        ('2.4', 'תקציב',
         '"התקציב החודשי שלי כ-8,000 שקל, בערך 2,200 דולר"',
         'המערכת שומרת את התקציב בדולרים', 'התקציב נשמר בשקלים ללא המרה'),
        ('2.5', 'אזור גיאוגרפי',
         '"אני מספקת לכל הארץ, בעיקר למרכז"',
         'המערכת מאשרת: טירגוט = ישראל, דגש על מרכז', 'טירגוט נשמר כ"עולמי"'),
        ('2.6', 'פלטפורמות',
         '"יש לי Google Ads ו-Facebook, TikTok לא"',
         'המערכת שומרת: Google + Meta בלבד', 'TikTok מופיע כפעיל'),
        ('2.7', 'קול המותג',
         '"המותג שלי חם, טבעי, כנה. אנחנו אוהבים משפחות שאוכלות טוב."',
         'המערכת מאשרת ושומרת את קול המותג', 'בפוסטים שנוצרים מאוחר יותר — טון תאגידי'),
        ('2.8', 'עמודי תוכן',
         '"אורח חיים אורגני, מתכונים עונתיים, סיפורי חווה, בריאות המשפחה"',
         'ארבעת הנושאים נשמרים כעמודי תוכן', 'נושאי ברירת מחדל גנריים'),
        ('2.9', 'הגבלות',
         '"אנא אל תפרסמו בשבת — מיום שישי 17:00 עד מוצ"ש"',
         'המערכת שומרת חוק שבת: שישי 17:00–שבת 23:00', 'ההגבלה לא נשמרת'),
        ('2.10', 'רמת סיכון',
         '"אני פתוחה לניסויים אבל לא לסיכון אגרסיבי"',
         'המערכת שומרת: רמת סיכון = בינונית', 'המערכת משתמשת בהגשה אגרסיבית'),
    ]
    for num, title, say, expect, report in topics:
        step_header(doc, num, title)
        action_line(doc, f'כתבי: {say}')
        para(doc, 'מה צריך לראות:', bold=True, size=10.5)
        expect_line(doc, expect)
        para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
        report_line(doc, report)
        checkbox(doc, 'בוצע ✓', indent=0.1)
        divider(doc)

    step_header(doc, '2.11', 'סיום השאלון — קבלת אסטרטגיה')
    action_line(doc, 'סיימי לענות על כל 10 הנושאים')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'הודעה: "האסטרטגיה שלך מוכנה"')
    expect_line(doc, 'סיכום: סוג עסק, מטרה, תקציב, טירגוט, מותג, פלטפורמות')
    expect_line(doc, 'כפתור ברור: "חברי את Google Ads"')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'אסטרטגיה לא נוצרת אחרי כל 10 נושאים')
    report_line(doc, 'הסיכום מציג ערכים שגויים')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 3 ─────────────────────────────────────────────────────────────
    phase_title(doc, 3, 'חיבור פלטפורמות פרסום', '20 דקות')

    step_header(doc, '3.1', 'חיבור Google Ads')
    action_line(doc, 'לחצי על "חיבור Google Ads"')
    action_line(doc, 'עברי את מסך ההרשאות של Google ← אשרי')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'נפתח מסך Google OAuth (של Google עצמה)')
    expect_line(doc, 'אחרי אישור — חזרה ל-Vigmis עם הודעה "Google Ads חובר"')
    expect_line(doc, 'שם חשבון ה-Google Ads שלך מופיע ברשימה')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'OAuth נכשל / מעביר לדף שגוי')
    report_line(doc, 'חשבון Google Ads לא מופיע ברשימה')
    report_line(doc, 'הסטטוס מציג "מנותק" אחרי חיבור מוצלח')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '3.2', 'חיבור Facebook + Instagram')
    action_line(doc, 'לחצי על "חיבור Facebook"')
    action_line(doc, 'עברי את OAuth של Meta ← אשרי את כל ההרשאות')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'מסך Meta OAuth נפתח עם בקשת הרשאות')
    expect_line(doc, 'הדף שלך ב-Facebook מופיע לבחירה')
    expect_line(doc, 'חשבון ה-Instagram שלך מופיע מקושר לדף')
    expect_line(doc, 'חשבון המודעות שלך מופיע לבחירה')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'OAuth נכשל')
    report_line(doc, 'הדף שלך לא מופיע')
    report_line(doc, 'Instagram לא מופיע מקושר לדף')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '3.3', 'בדיקת Dashboard אחרי חיבור')
    action_line(doc, 'חזרי לדשבורד הראשי')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'Google Ads מסומן "מחובר" (ירוק)')
    expect_line(doc, 'Facebook מסומן "מחובר" (ירוק)')
    expect_line(doc, 'Instagram מסומן "מחובר" (ירוק)')
    expect_line(doc, 'שמות החשבונות קריאים — לא מספרים בלבד')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'פלטפורמות מציגות "מנותק" אחרי חיבור')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 4 ─────────────────────────────────────────────────────────────
    phase_title(doc, 4, 'קמפיינים', '30 דקות')

    step_header(doc, '4.1', 'יצירת קמפיין ראשון')
    note_box(doc, 'Vigmis אמורה ליצור קמפיין אוטומטית אחרי אישור האסטרטגיה')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'קמפיין מופיע ברשימת הקמפיינים')
    expect_line(doc, 'שם בפורמט: VIGMIS_GOOGLE_SEARCH_2026-...')
    expect_line(doc, 'תקציב יומי: ~$73 (מתוך $2,200 חודשי)')
    expect_line(doc, 'סטטוס: "ממתין לאישור" ← "פעיל" אחרי אישור')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'לא נוצר קמפיין אחרי אישור האסטרטגיה')
    report_line(doc, 'תקציב שגוי (תקציב חודשי שלם במקום יומי)')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '4.2', 'השהיית קמפיין דרך הצ׳אט')
    action_line(doc, 'פתחי את הצ׳אט עם Vigmis')
    action_line(doc, 'כתבי: "עצרי את הקמפיין של Google Ads להיום"')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'Vigmis עונה: "אשהה את הקמפיין עכשיו"')
    expect_line(doc, 'סטטוס הקמפיין משתנה ל-"מושהה"')
    expect_line(doc, 'אישור בצ׳אט שהפעולה בוצעה')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'Vigmis עונה אבל לא משהה בפועל')
    report_line(doc, 'סטטוס לא משתנה')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '4.3', 'חידוש קמפיין')
    action_line(doc, 'כתבי בצ׳אט: "הפעילי חזרה את הקמפיין"')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'סטטוס הקמפיין משתנה ל-"פעיל"')
    expect_line(doc, 'אישור בצ׳אט')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '4.4', 'נתוני ביצועים (אחרי 3 ימים לפחות)')
    action_line(doc, 'פתחי את ה-Dashboard אחרי 3 ימי פעילות')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'נוצגות חשיפות, קליקים, הוצאה, המרות')
    expect_line(doc, 'ROAS מחושב (הכנסה חלקי הוצאה)')
    expect_line(doc, 'גרף מגמה יומי')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'כל המדדים מציגים 0 אחרי 3+ ימים')
    report_line(doc, 'ROAS חסר')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 5 ─────────────────────────────────────────────────────────────
    phase_title(doc, 5, 'סושיאל מדיה — פוסטים', '45 דקות')

    step_header(doc, '5.1', 'פוסטים שבועיים נוצרים')
    note_box(doc, 'Vigmis יוצרת 7 פוסטים טיוטה בכל שבוע. ניתן גם להפעיל ידנית.')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, '7 פוסטים בסטטוס "טיוטה" בתיבת הנכנסים')
    expect_line(doc, 'הפוסטים בעברית — לא באנגלית!')
    expect_line(doc, 'חלק לפייסבוק + חלק לאינסטגרם')
    expect_line(doc, 'הטון חם ומשפחתי (לא תאגידי)')
    expect_line(doc, 'הפוסטים קשורים לאורגני / מזון / בריאות')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'פוסטים באנגלית — חמור!')
    report_line(doc, 'פחות מ-7 פוסטים')
    report_line(doc, 'כל הפוסטים לאותה פלטפורמה')
    report_line(doc, 'תוכן גנרי שלא קשור לעסק')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '5.2', 'אישור פוסט')
    action_line(doc, 'פתחי פוסט ← בדקי את התוכן ← לחצי "אשרי"')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'סטטוס הפוסט משתנה ל-"מאושר"')
    expect_line(doc, 'הפוסט עובר לתור הפרסום')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'לחצן "אשרי" לא קיים')
    report_line(doc, 'הסטטוס לא משתנה')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '5.3', 'עריכת פוסט לפני אישור')
    action_line(doc, 'לחצי "ערכי" על פוסט ← שנה משפט אחד ← אשרי')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'השינוי נשמר')
    expect_line(doc, 'הפוסט המאושר מכיל את הטקסט הערוך')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'הטקסט המקורי חוזר אחרי שמירה')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '5.4', 'דחיית פוסט')
    action_line(doc, 'לחצי "דחי" על פוסט')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'הפוסט מסומן "נדחה"')
    expect_line(doc, 'הפוסט לא יפורסם')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'פוסט שנדחה מתפרסם למחרת')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '5.5', 'אישור פרסום בפועל')
    action_line(doc, 'בדקי את דף הפייסבוק / אינסטגרם שלך בזמן המתוזמן')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'הפוסט מופיע בדף Facebook שלך')
    expect_line(doc, 'הפוסט מופיע בחשבון Instagram שלך')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'הפוסט לא פורסם בזמן')
    report_line(doc, 'פורסם לחשבון הלא נכון')
    report_line(doc, 'פוסט כפול (פורסם פעמיים)')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 6 ─────────────────────────────────────────────────────────────
    phase_title(doc, 6, 'תגובות — תיבת הנכנסים', '45 דקות')

    step_header(doc, '6.1', 'תגובות מגיעות ומסווגות')
    action_line(doc, 'פתחי את "תיבת התגובות" / Comments Inbox')
    para(doc, 'מה צריך לראות לכל תגובה:', bold=True, size=10.5)
    expect_line(doc, 'שם הכותב')
    expect_line(doc, 'טקסט התגובה בשפה המקורית')
    expect_line(doc, 'סימון פלטפורמה: FB / IG')
    expect_line(doc, 'תג סנטימנט: חיובי / שאלה / תלונה / ספאם / אחר')
    expect_line(doc, 'ציון עדיפות: גבוה / בינוני / נמוך')
    expect_line(doc, 'תשובה מוצעת בעברית בסגנון המותג')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'אין תגובות אפילו שיש פעילות אמיתית בדפים')
    report_line(doc, 'תשובות מוצעות באנגלית')
    report_line(doc, 'כל התגובות מסווגות לאותה קטגוריה')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '6.2', 'מענה על שאלת לקוח')
    action_line(doc, 'מצאי תגובה עם שאלה (למשל: "האם החמאת שקדים מתאימה לסוכרתיים?")')
    action_line(doc, 'בדקי את התשובה המוצעת ← לחצי "שלחי תשובה"')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'התשובה נשלחת לתגובה המקורית ב-Facebook/Instagram')
    expect_line(doc, 'סטטוס התגובה משתנה ל-"נענה"')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'לחיצה על "שלח" לא שולחת בפועל לפלטפורמה')
    report_line(doc, 'התשובה נשלחת עם טקסט שונה ממה שנכתב')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '6.3', 'התעלמות מספאם')
    action_line(doc, 'מצאי תגובת ספאם ← לחצי "התעלמי"')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'התגובה נעלמת מהתיבה')
    expect_line(doc, 'לא מופיעה שוב בסנכרון הבא')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '6.4', 'זיהוי משבר — בדיקה קריטית!')
    note_box(doc, 'זו בדיקה חשובה מאוד — ויגמיס חייבת לזהות תגובות קריטיות באופן מיידי')
    action_line(doc, 'כתבי תגובת בדיקה (בפייסבוק בעצמך): "מצאתי עובש במוצר! זה מסוכן!"')
    para(doc, 'מה צריך לראות תוך שעה:', bold=True, size=10.5)
    expect_line(doc, 'התגובה מופיעה עם תג אדום "משבר"')
    expect_line(doc, 'עדיפות = גבוהה מאוד')
    expect_line(doc, 'מגיע מייל התראה')
    expect_line(doc, 'ה-Dashboard מציג באנר אזהרה')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'התגובה מסווגת רק כ"תלונה" ללא תג משבר')
    report_line(doc, 'לא מגיע מייל התראה')
    report_line(doc, 'אין באנר אזהרה ב-Dashboard')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 7 ─────────────────────────────────────────────────────────────
    phase_title(doc, 7, 'צ׳אט עם היועץ AI', '30 דקות')

    convos = [
        ('7.1', 'בקשת סיכום ביצועים',
         '"איך הקמפיינים שלי עברו השבוע?"',
         ['סיכום עם מספרים אמיתיים (חשיפות, קליקים, הוצאה, ROAS)',
          'השוואה לשבוע קודם',
          'המלצה קונקרטית'],
         ['תשובה כללית ללא נתונים אמיתיים',
          'המלצה מסוג "שקולי להגדיל"']),
        ('7.2', 'בקשת הסבר',
         '"למה ויגמיס השהתה את הקמפיין אתמול?"',
         ['סיבה ספציפית (למשל: "עלות קליק עלתה 40% מעל הסף")',
          'זמן ותאריך הפעולה',
          'מה לעשות עכשיו'],
         ['תשובה גנרית "לא יודע"',
          'ללא תאריך / ללא סיבה']),
        ('7.3', 'בקשת יצירת תוכן',
         '"צרי פוסט על קציר הדבש שלנו באביב"',
         ['פוסט בעברית',
          'בסגנון חם ומשפחתי',
          'עם האשטגים בעברית ואנגלית',
          'אפשרות לאשר/לערוך מהצ׳אט'],
         ['פוסט באנגלית', 'סגנון תאגידי']),
        ('7.4', 'בקשת המלצה תקציבית',
         '"המכירות עלו 30% השבוע, האם להגדיל תקציב?"',
         ['ניתוח ביצועים נוכחיים',
          'המלצה קונקרטית עם סכום מוצע',
          'הצעה לבצע את השינוי'],
         ['תמיד אומר "כן" ללא ניתוח',
          'תמיד אומר "התייעצי עם יועץ"']),
    ]
    for num, title, say, expects, reports in convos:
        step_header(doc, num, title)
        action_line(doc, f'כתבי בצ׳אט: {say}')
        para(doc, 'מה צריך לראות:', bold=True, size=10.5)
        for e in expects: expect_line(doc, e)
        para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
        for r in reports: report_line(doc, r)
        checkbox(doc, 'בוצע ✓', indent=0.1)
        divider(doc)

    step_header(doc, '7.5', 'בדיקת מכסת הודעות (תוכנית חינם)')
    action_line(doc, 'שלחי 50+ הודעות בחודש (תוכנית חינם)')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'אזהרה לפני הגעה למגבלה: "נשאר לך X שיחות החודש"')
    expect_line(doc, 'כשמגיעים למגבלה — תשובות מקוצרות (Degrade mode)')
    expect_line(doc, 'הצעת שדרוג ל-Pro')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'אין אזהרה לפני המגבלה')
    report_line(doc, 'חסימה מוחלטת ללא הסבר')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 8 ─────────────────────────────────────────────────────────────
    phase_title(doc, 8, 'חיוב ותוכניות', '20 דקות')

    step_header(doc, '8.1', 'צפייה בשימוש הנוכחי')
    action_line(doc, 'עברי לעמוד Billing')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'תוכנית: Free')
    expect_line(doc, 'מספר שיחות AI שנוצלו החודש')
    expect_line(doc, 'כמות תגובות שטופלו')
    expect_line(doc, 'הוצאת פרסום מנוהלת')
    expect_line(doc, 'עמודות התקדמות לעבר המגבלה')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'שימוש מציג 0 למרות פעילות רבה')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '8.2', 'שדרוג ל-Pro')
    action_line(doc, 'לחצי "שדרגי ל-Pro" ← Paddle checkout נפתח')
    action_line(doc, 'בצעי תשלום בדיקה (כרטיס טסט: 4242 4242 4242 4242)')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'Paddle checkout נפתח עם המחיר הנכון ($49/חודש)')
    expect_line(doc, 'אחרי תשלום — תוכנית משתנה ל-"Pro"')
    expect_line(doc, 'מגבלות גבוהות יותר מופעלות')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'Checkout לא נפתח')
    report_line(doc, 'אחרי תשלום — עדיין מוצגת תוכנית Free')
    report_line(doc, 'מחיר שגוי')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 9 ─────────────────────────────────────────────────────────────
    phase_title(doc, 9, 'דוחות וניתוחים', '20 דקות')

    step_header(doc, '9.1', 'דוח יומי במייל')
    action_line(doc, 'בדקי את תיבת הדואר בבוקר (אחרי לילה ראשון פעיל)')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'מייל: "הדוח היומי שלך מ-Vigmis"')
    expect_line(doc, 'כולל: הוצאה, ROAS, קמפיין מוביל')
    expect_line(doc, 'פעולות שבוצעו אוטומטית')
    expect_line(doc, 'אישורים ממתינים (אם יש)')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'לא מגיע מייל')
    report_line(doc, 'מייל ריק ללא נתונים')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '9.2', 'Analytics ב-Dashboard')
    action_line(doc, 'פתחי את עמוד ה-Analytics')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'גרף ביצועים לפי יום')
    expect_line(doc, 'פירוט לפי Google / Meta')
    expect_line(doc, 'פילוח גיאוגרפי (מרכז ישראל בראש)')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'גרפים לא מוצגים (ריבועים לבנים)')
    report_line(doc, 'פילוח גיאוגרפי מציג "תל אביב" בלבד')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '9.3', 'ייצוא נתונים')
    action_line(doc, 'לחצי "ייצוא" / Export ← הורידי CSV')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'קובץ CSV מוריד תוך 10 שניות')
    expect_line(doc, 'הקובץ מכיל: קמפיין, הוצאה, המרות, ROAS')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'כפתור ייצוא לא קיים')
    report_line(doc, 'CSV ריק')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── PHASE 10 ─────────────────────────────────────────────────────────────
    phase_title(doc, 10, 'מקרי קצה', '30 דקות')

    step_header(doc, '10.1', 'ציות לשבת')
    action_line(doc, 'בדקי שבין שישי 17:00 לשבת 23:00 — קמפיינים מושהים')
    action_line(doc, 'בדקי שאחרי מוצ"ש — קמפיינים חוזרים לפעול')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'קמפיינים ממשיכים לרוץ בשבת')
    report_line(doc, 'אזור זמן שגוי (לא UTC+3 / ישראל)')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '10.2', 'התראת תקציב נמוך')
    note_box(doc, 'ניתן לסמולציה על ידי הפחתת התקציב היומי ל-$5 באמצע יום')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'התראה ב-Dashboard: "התקציב עומד להסתיים"')
    expect_line(doc, 'הודעת צ׳אט מ-Vigmis')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '10.3', 'ROAS גבוה — הצעת הגדלה')
    note_box(doc, 'אחרי 3 ימים רצופים עם ROAS מעל 8x — Vigmis אמורה להציע להגדיל')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'הודעה יזומה מ-Vigmis: "ה-ROAS שלך הוא 8x — שקלי להגדיל תקציב"')
    expect_line(doc, 'המלצה עם סכום קונקרטי')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'אין הודעה יזומה')
    report_line(doc, 'הודעה מגיעה רק ביום 7+')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    divider(doc)
    step_header(doc, '10.4', 'ניתוק וחיבור מחדש של Meta')
    action_line(doc, 'נתקי Facebook ← חברי מחדש')
    para(doc, 'מה צריך לראות:', bold=True, size=10.5)
    expect_line(doc, 'אחרי חיבור מחדש — אותו דף נבחר')
    expect_line(doc, 'תיבת התגובות לא התאפסה')
    expect_line(doc, 'הפוסטים בתור לא נמחקו')
    para(doc, 'מה לדווח:', bold=True, size=10.5, color=C_RED)
    report_line(doc, 'כל הפוסטים נמחקו בניתוק')
    report_line(doc, 'תיבת התגובות התאפסה')
    checkbox(doc, 'בוצע ✓', indent=0.1)

    add_page_break(doc)

    # ── BUG REPORT TEMPLATE ──────────────────────────────────────────────────
    bug_report_template(doc)

    doc.add_paragraph()

    # ── SEVERITY TABLE ───────────────────────────────────────────────────────
    heading(doc, 'הגדרות חומרת תקלה', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdrs = ['חומרה', 'הגדרה', 'דוגמאות']
    for i, h in enumerate(hdrs):
        set_cell_bg(table.rows[0].cells[i], '1a213e')
        p = table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h); r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        set_rtl(p)
    data = [
        ('קריטי', 'פיצ׳ר שבור לחלוטין', 'לא ניתן להירשם, קמפיין לא נעצר'),
        ('גבוה', 'פיצ׳ר שבור, יש עקיפה', 'שפת פוסטים שגויה, חיוב שגוי'),
        ('בינוני', 'התנהגות שגויה חלקית', 'מדד לא מוצג, סדר רשימה שגוי'),
        ('נמוך', 'בעיה ויזואלית / UX', 'כתיב, יישור, צבע'),
    ]
    for i, (sev, defn, ex) in enumerate(data):
        row = table.rows[i + 1]
        bg = ['D93025', 'E37A00', '1a73e8', '0d904f'][i]
        set_cell_bg(row.cells[0], bg)
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(sev); r0.font.bold = True; r0.font.color.rgb = RGBColor(255,255,255); r0.font.size = Pt(10)
        set_rtl(p0)
        for ci, txt in enumerate([defn, ex], 1):
            p = row.cells[ci].paragraphs[0]
            r = p.add_run(txt); r.font.size = Pt(10)
            set_rtl(p)

    doc.add_paragraph()
    divider(doc)
    p = doc.add_paragraph('בהצלחה! 🌿 — צוות Vigmis')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = C_NAVY
        p.runs[0].font.bold = True
    set_rtl(p)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")

if __name__ == '__main__':
    build()
