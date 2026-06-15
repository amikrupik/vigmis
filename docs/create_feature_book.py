#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Creates VIGMIS_FEATURES_HE.docx — a comprehensive Hebrew product feature book.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"C:\vigmis\vigmis-main\docs\VIGMIS_FEATURES_HE.docx"

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_rtl(paragraph):
    """Force right-to-left paragraph direction."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def set_ltr(paragraph):
    """Force left-to-right paragraph direction (for LTR content inside RTL doc)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '0')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)

def add_heading(doc, text, level=1, rtl=True):
    p = doc.add_heading(text, level=level)
    if rtl:
        set_rtl(p)
    return p

def add_paragraph(doc, text, bold=False, rtl=True, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if rtl:
        set_rtl(p)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, rtl=True):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    if rtl:
        set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bold_paragraph(doc, label, text, rtl=True):
    """Add a paragraph with a bold label followed by normal text."""
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r2 = p.add_run(text)
    if rtl:
        set_rtl(p)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with Hebrew headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '3949AB')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            # Alternate row shading
            if r_idx % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8EAF6')
                tcPr.append(shd)

    return table

def add_page_break(doc):
    doc.add_page_break()

def add_info_box(doc, text, rtl=True):
    """Add a styled info/note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    if rtl:
        set_rtl(p)
    return p

# ── Document creation ──────────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title Page ────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    run = p_title.add_run('VIGMIS')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_subtitle = doc.add_paragraph()
    run2 = p_subtitle.add_run('מנהל הפרסום שלך עובד 24/7')
    run2.bold = True
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0x30, 0x49, 0xAB)
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p_subtitle)

    p_desc = doc.add_paragraph()
    run3 = p_desc.add_run('ספר פיצ\'רים מקיף — עברית')
    run3.font.size = Pt(13)
    run3.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p_desc)

    p_date = doc.add_paragraph()
    run4 = p_date.add_run('יוני 2026')
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 1: מהו Vigmis
    # ══════════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'פרק 1: מהו Vigmis', level=1)

    add_paragraph(doc,
        'Vigmis הוא מנהל פרסום מבוסס בינה מלאכותית לעסקים קטנים ובינוניים. '
        'במקום לשכור סוכנות שיווק, לעבוד עם פרילנסרים, או ללמוד בעצמך איך מנהלים '
        'קמפיינים ב-Google, Meta ו-TikTok — Vigmis עושה את הכל עבורך: '
        'מנהל את הקמפיינים, יוצר תוכן ופרסומות, מנתח נתונים, מקבל החלטות בזמן אמת, '
        'ומדווח לך בשפה פשוטה מה קורה ולמה. '
        'בעל העסק מתמקד בעסק שלו — Vigmis מתמקד בפרסום.'
    )

    add_paragraph(doc,
        'Vigmis מחובר ישירות לחשבונות הפרסום שלך ב-Google Ads, Meta (פייסבוק ואינסטגרם), '
        'ו-TikTok. הוא גם מתחבר ל-Shopify, Google Analytics 4, ומערכות נוספות — '
        'כדי לדעת לא רק כמה הוצאת, אלא כמה הרווחת באמת.'
    )

    add_bold_paragraph(doc, 'למי מיועד Vigmis? ',
        'לכל בעל עסק שרוצה להופיע לפני הלקוחות הנכונים ברשת, '
        'בלי צורך בידע טכני, בלי לנהל מספר ספקים, ובלי לבזבז שעות על דוחות. '
        'מגנרת ביגוד ועד קליניקת שיניים, ממסעדה ועד חנות אונליין.')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 2: כניסה ראשונה — Onboarding
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 2: כניסה ראשונה — האשף הראשוני', level=1)

    add_paragraph(doc,
        'כשנכנסים לראשונה ל-Vigmis, אשף חכם לוקח אותך תוך כ-10 דקות '
        'מהרשמה ועד קמפיין פעיל. לא צריך ידע טכני — רק לענות על כמה שאלות פשוטות.'
    )

    add_heading(doc, 'שלב 1: חיבור פלטפורמות הפרסום', level=2)
    add_paragraph(doc,
        'הצעד הראשון הוא לחבר את חשבונות הפרסום שלך. '
        'Vigmis תומך בכניסה חד-פעמית (OAuth) — ללא שמירת סיסמאות:'
    )
    add_bullet(doc, 'Google Ads — לחיצה אחת, בחירת חשבון הפרסום, וסיום')
    add_bullet(doc, 'Meta (פייסבוק + אינסטגרם) — כולל בחירת הדף העסקי וחשבון הפרסום')
    add_bullet(doc, 'TikTok Ads — בקרוב (ממתין לאישור App Review)')

    add_paragraph(doc,
        'אפשר לחבר פלטפורמה אחת ולהמשיך — לא חייבים את כולן מההתחלה. '
        'Vigmis יציין אם יש פלטפורמות שמומלץ להוסיף בהמשך על בסיס האסטרטגיה שלך.',
        rtl=True
    )

    add_heading(doc, 'שלב 2: שאלון AI — ריאיון עסקי', level=2)
    add_paragraph(doc,
        'אחרי החיבור, צ\'אט AI מנהל איתך ריאיון קצר. '
        'השאלות מכסות את הדברים הבסיסיים שכל מנהל פרסום צריך לדעת:'
    )
    add_bullet(doc, 'כתובת האתר שלך')
    add_bullet(doc, 'מה המטרה — לידים, מכירות, הגברת מודעות')
    add_bullet(doc, 'תקציב חודשי ואחוז שתרצה ש-Vigmis ינהל')
    add_bullet(doc, 'אזור גיאוגרפי — ישראל, כל העולם, אזור ספציפי')
    add_bullet(doc, 'דברים שחשוב לך לא לפרסם')
    add_bullet(doc, 'הערות חופשיות — כל מה שחשוב לך')

    add_info_box(doc,
        'הצ\'אט מזהה אוטומטית את שפתך (עברית, אנגלית, ערבית ו-10 שפות נוספות) '
        'ומנהל את השיחה בשפה שבה אתה מדבר.'
    )

    add_heading(doc, 'שלב 3: ניתוח האתר ומחקר שוק', level=2)
    add_paragraph(doc,
        'לאחר השאלון, Vigmis מבצע שלושה דברים במקביל:'
    )
    add_bullet(doc, 'סורק את האתר שלך ומבין מה אתה מוכר, למי, ובאיזה מחיר')
    add_bullet(doc, 'מחקר שוק בזמן אמת (Perplexity Sonar Pro) — מה עושים המתחרים, מה מגמות החיפוש, נתוני benchmark בתעשייה שלך')
    add_bullet(doc, 'בונה תוכנית אסטרטגית מלאה בהתבסס על כל הנתונים')

    add_paragraph(doc,
        'אם האתר שלך בנוי בצורה שקשה לסרוק (React/Next.js ללא SSR), '
        'Vigmis ישאל אותך לתאר את העסק במילים שלך ויבנה את האסטרטגיה על בסיס זה.',
        rtl=True
    )

    add_heading(doc, 'שלב 4: הצגת תוכנית האסטרטגיה', level=2)
    add_paragraph(doc,
        'Vigmis מציג בפניך תוכנית שלמה הכוללת:'
    )
    add_bullet(doc, 'נרטיב אסטרטגי — מה הגישה, למה, ואיך מודדים הצלחה')
    add_bullet(doc, 'פיצול תקציב — כמה ל-Google, כמה ל-Meta, ולמה')
    add_bullet(doc, 'קהל יעד מוגדר')
    add_bullet(doc, 'הערכת CPC (עלות לקליק) בתעשייה שלך')
    add_bullet(doc, 'אומדן קליקים, לידים ונקודת איזון')
    add_bullet(doc, 'המלצות לתוכן קריאטיבי (מה לעשות, לאיזה פלטפורמה, בכמה)')
    add_bullet(doc, 'אזהרות תקציב (אם התקציב נמוך מדי לשוק)')

    add_paragraph(doc,
        'אתה יכול לבקש שינויים בתוכנית, ו-Vigmis יסביר למה הוא ממליץ אחרת '
        '— ואם תחליט לעקוף את ההמלצה, הוא יכבד את ההחלטה שלך.',
        rtl=True
    )

    add_heading(doc, 'שלב 5: בחירת סוג קריאייטיב ראשון', level=2)
    add_paragraph(doc,
        'תבחר איזה סוג פרסומת תרצה לייצר ראשון: '
        'סרטון אווטאר, סרטון קולנועי, אנימציה, תמונה ממוחשבת, '
        'או להתחיל בלי קריאייטיב ולהוסיף מאוחר יותר.'
    )

    add_heading(doc, 'שלב 6: הגדרת מעקב המרות (Tracking)', level=2)
    add_paragraph(doc,
        'Vigmis מספק קוד מעקב קצר שמוטמע באתרך. '
        'הקוד מעקב אחרי פניות, רכישות, ופעולות ערך. '
        'אפשר גם לחבר חנות Shopify לקבלת נתונים מדויקים על הזמנות בפועל.'
    )

    add_heading(doc, 'שלב 7: השקת הקמפיינים', level=2)
    add_paragraph(doc,
        'לאחר אישור התוכנית, לחיצה אחת משיקה את כל הקמפיינים. '
        'Vigmis יוצר את מבנה הקמפיינים ב-Google ו-Meta, '
        'מגדיר את קהלי היעד, מקצה תקציבים — הכל אוטומטי.'
    )

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 3: הדשבורד
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 3: הדשבורד — המרכז הניהולי', level=1)

    add_paragraph(doc,
        'הדשבורד הוא המקום שבו אתה רואה הכל ושולט בהכל. '
        'הוא מחולק ל-11 לשוניות (טאבים), כל אחת עם מטרה ברורה. '
        'ניתן לגשת לכל לשונית ישירות מהסרגל הצדדי.'
    )

    # Tab 1: Overview
    add_heading(doc, 'לשונית 1: סקירה כללית (Overview)', level=2)
    add_paragraph(doc,
        'זו הדף הראשי שתראה כשנכנסים לדשבורד. '
        'הוא מציג את תמונת המצב השלמה של הפרסום שלך בלמבוט אחד:'
    )
    add_bullet(doc, 'מדדי ביצועים: הוצאה כוללת, ROAS (תשואה על הוצאת פרסום), קונברסיות, CPA (עלות לרכישה), CTR (אחוז הקלקה), קליקים וחשיפות')
    add_bullet(doc, 'השוואה לתקופה קודמת — האם הביצועים השתפרו או הידרדרו')
    add_bullet(doc, 'ציון נראות ב-AI (GEO) — עד כמה הולך מוצא אותך כשמחפשים בצ\'אט AI')
    add_bullet(doc, 'בעיות דחופות ואזהרות — כמות הבעיות הפתוחות שדורשות תשומת לב')
    add_bullet(doc, 'Conversion Intelligence — השוואת הזמנות Shopify אמיתיות מול דיווחי הפלטפורמות (ROAS אמיתי מול ROAS מוצהר)')
    add_bullet(doc, 'פעולות AI אחרונות — מה עשה Vigmis ב-24 השעות האחרונות')
    add_bullet(doc, 'קמפיינים פעילים ותקציב יומי כולל')
    add_bullet(doc, 'הצעד הבא — המלצות מהירות ("צור פוסט שבועי", "צור קריאייטיב")')
    add_bullet(doc, 'כפתור עצירה חירום — עוצר את כל הקמפיינים ברגע')

    # Tab 2: Strategy
    add_heading(doc, 'לשונית 2: אסטרטגיה (Strategy)', level=2)
    add_paragraph(doc,
        'כאן תמצא את התוכנית האסטרטגית המלאה של Vigmis לעסק שלך, '
        'כולל:'
    )
    add_bullet(doc, 'הנרטיב האסטרטגי — הסבר מלא של הגישה')
    add_bullet(doc, 'תובנות שוק — מה קורה בתעשייה שלך עכשיו')
    add_bullet(doc, 'קהל יעד מוגדר (ICP)')
    add_bullet(doc, 'הקצאת תקציב לפי פלטפורמה')
    add_bullet(doc, 'ניתוח תקציב — מינימום להיכנס, מומלץ ללמידה, מומלץ לשוטף')
    add_bullet(doc, 'ציוני ביטחון (Confidence Scores) — עד כמה Vigmis בטוח בהמלצות')
    add_bullet(doc, 'גורמי סיכון — 5 הסיכונות הגדולים, ההסתברות שיקרו, ומה לעשות')
    add_bullet(doc, 'נגד-טיעון (Counter-argument) — "למה לא ללכת על X?"')
    add_bullet(doc, 'סטטיסטיקות עם מקורות (מ-WordStream, Meta, Google 2025)')
    add_bullet(doc, 'ציון DNA מותג — צבעים, פונטים, אלמנטים שלא לשנות')

    add_bold_paragraph(doc, 'פעולה זמינה: ',
        'לחיצה על "הרץ ניתוח מחדש" מפעילה מחקר שוק חדש ומעדכנת את האסטרטגיה.')

    # Tab 3: Analytics
    add_heading(doc, 'לשונית 3: אנליטיקה (Analytics)', level=2)
    add_paragraph(doc,
        'מציג נתוני ביצועים מפורטים לפי תקופה (7 / 30 / 90 ימים):'
    )
    add_bullet(doc, 'הוצאה כוללת, ערך קונברסיות, מספר קונברסיות')
    add_bullet(doc, 'ROAS, CPA, CTR, קליקים, חשיפות')
    add_bullet(doc, 'השוואה לתקופה קודמת (% שינוי)')
    add_bullet(doc, 'גרפים יומיים — תנועת ביצועים לאורך הזמן')
    add_bullet(doc, 'ויזואליזציה של משפך קונברסיה')
    add_bullet(doc, 'ייצוא לאקסל (CSV) ול-HTML מוכן להדפסה')

    # Tab 4: Campaigns
    add_heading(doc, 'לשונית 4: קמפיינים (Campaigns)', level=2)
    add_paragraph(doc,
        'רשימת כל הקמפיינים הפעילים עם מידע מלא:'
    )
    add_bullet(doc, 'פלטפורמה (Google/Meta/TikTok), סוג קמפיין, סטטוס')
    add_bullet(doc, 'תקציב יומי בדולר')
    add_bullet(doc, 'ציון תקופת הלמידה — Vigmis מסביר מתי הקמפיין ייצא מתקופת הלמידה')
    add_bullet(doc, 'הודעות שגיאה (אם יש)')
    add_bullet(doc, 'השהיה/חידוש קמפיין בודד')
    add_bullet(doc, 'עריכת תקציב ישירות מהטבלה')
    add_bullet(doc, 'השהיה/חידוש כל הקמפיינים בבת אחת')

    add_info_box(doc,
        'שם הקמפיין נוצר אוטומטית בפורמט: VIGMIS_PLATFORM_TYPE_DATE. '
        'לפני השקה, Vigmis בודק שהאתר שלך עומד בדרישות בסיסיות של דף נחיתה — '
        'כדי לא לבזבז תקציב על קמפיין שלא ימיר.'
    )

    # Tab 5: Creative
    add_heading(doc, 'לשונית 5: קריאייטיב (Creative)', level=2)
    add_paragraph(doc,
        'מרכז ייצור התוכן של Vigmis. כאן יוצרים את החומרים הפרסומיים — '
        'סרטונים, תמונות, טקסטים.'
    )

    add_heading(doc, 'המלצות AI', level=3)
    add_paragraph(doc,
        '3 קונספטים מוכנים מ-Creative Director AI, מבוססים על האסטרטגיה שלך. '
        'לחיצה על "השתמש" ממלאת את הטופס אוטומטית.'
    )

    add_heading(doc, 'ייצור סרטונים', level=3)
    table_video = add_table(doc,
        ['סוג', 'טכנולוגיה', 'מחיר', 'רזולוציה', 'מתאים ל'],
        [
            ['סרטון אווטאר', 'HeyGen', '$15', '1080p', 'מצגת מוצר עם דמות AI'],
            ['סרטון קולנועי', 'Replicate', '$12', '1080p 16:9', 'פרסומות YouTube/Meta'],
            ['סרטון אנימציה', 'Replicate', '$8', '720p 9:16', 'ריל, TikTok, Stories'],
        ]
    )

    doc.add_paragraph()

    add_heading(doc, 'תמונות ועיצוב', level=3)
    add_bullet(doc, 'יצירת תמונות AI (gpt-image-1) — $5 לתמונה')
    add_bullet(doc, 'Best-of-3: Vigmis מייצר 3 תמונות, AI Critic בוחר את הטובה ביותר')
    add_bullet(doc, 'אם הציון נמוך מ-75/100, Vigmis מנסה שוב אוטומטית')

    add_heading(doc, 'Ad Copy — כתיבת מודעות', level=3)
    add_bullet(doc, '6 וריאציות טקסט מותאמות לפלטפורמה')
    add_bullet(doc, 'עמידה במגבלות תווים (Google, Meta, TikTok)')
    add_bullet(doc, 'ציון חיזוי לכל וריאציה')
    add_bullet(doc, 'כפתור "השתמש לסרטון" — ממלא ברייף ייצור הסרטון')

    add_heading(doc, 'ספריית קריאייטיב', level=3)
    add_paragraph(doc,
        'כל הקריאייטיבים שנוצרו מסודרים בגלריה עם סטטוס, '
        'אפשרויות הורדה, צפייה ואישור.'
    )

    add_heading(doc, 'ספריית נכסי מותג', level=3)
    add_bullet(doc, 'העלאת לוגו, תמונות מוצר, ושאר חומרי המותג')
    add_bullet(doc, 'Vigmis מזריק את הלוגו לכל סרטון שנוצר — אוטומטי')
    add_bullet(doc, 'פורמטים נתמכים: JPG, PNG, GIF, WebP, MP4, MOV (עד 10MB)')

    add_heading(doc, 'מערכת רוויזיות', level=3)
    add_table(doc,
        ['רוויזיה', 'מחיר'],
        [
            ['1-3 (רוויזיות 0-2)', 'חינם'],
            ['4-6 (רוויזיות 3-5)', '50% מהמחיר המקורי'],
            ['7+ (רוויזיה 6 ומעלה)', 'חסום — יש להתחיל קריאייטיב חדש'],
        ]
    )
    doc.add_paragraph()
    add_paragraph(doc,
        'בכל רוויזיה אפשר לציין "מה לשמור" ו"מה לשנות", '
        'כדי שVigmis לא "ישבור" דברים שעבדו.'
    )

    # Tab 6: Intelligence
    add_heading(doc, 'לשונית 6: מודיעין (Intelligence)', level=2)
    add_paragraph(doc,
        'כלי המחקר והאופטימיזציה המתקדמים של Vigmis. '
        'נחלק לחמישה תחומים:'
    )

    add_heading(doc, 'A/B Testing', level=3)
    add_bullet(doc, 'יצירת ניסוי A/B — בחירת פלטפורמה, מטרה, וריאנטים A ו-B')
    add_bullet(doc, 'Vigmis מציע המלצת ניסוי מוכנה לפני הטופס (מבוסס אסטרטגיה)')
    add_bullet(doc, 'סיום אוטומטי: Z-test סטטיסטי, מינימום 50 קליקים לצד, 7 ימים')
    add_bullet(doc, 'כפיית סיום אחרי 30 יום מניסוי שלא הסתיים')
    add_bullet(doc, 'הפסקה אוטומטית של הוריאנט המפסיד ב-Meta')
    add_bullet(doc, 'כתיבת Decision Protocol מפורט (Vigmis מסביר את ההחלטה)')

    add_heading(doc, 'גילוי קהלים (Audience Discovery)', level=3)
    add_bullet(doc, 'AI מזהה פלחי קהל רווחיים שעדיין לא מכוסים')
    add_bullet(doc, '"הוסף לקמפיין" — מעביר לעריכת הקמפיין')

    add_heading(doc, 'ניתוח מתחרים (Competitor Analysis)', level=3)
    add_bullet(doc, 'מחובר ל-Facebook Ad Library')
    add_bullet(doc, 'מציג פרסומות פעילות של מתחרים כשMeta מחובר')

    add_heading(doc, 'Budget Shift — העברת תקציב חכמה', level=3)
    add_bullet(doc, 'המלצה להעביר תקציב מקמפיין חלש לחזק')
    add_bullet(doc, 'הסבר מפורט ולחצן אישור')

    add_heading(doc, 'CRO Audit — ניתוח דף הנחיתה', level=3)
    add_bullet(doc, 'Vigmis בודק את דף הנחיתה שלך ומגלה חסמי המרה')
    add_bullet(doc, 'ציון מוכנות (0-100) + רשימת תיקונים ממוינת לפי דחיפות')

    add_heading(doc, 'ניתוח אלמנטים קריאטיביים', level=3)
    add_bullet(doc, 'Vigmis מזהה אילו נושאים, הוקים וסגנונות מביאים הכי הרבה קליקים')
    add_bullet(doc, 'הנתונים מוזרמים חזרה לCreative Director לשיפור הקריאטיבים הבאים')

    # Tab 7: GEO
    add_heading(doc, 'לשונית 7: נראות ב-AI / GEO', level=2)
    add_paragraph(doc,
        'כשמישהו שואל את ChatGPT, Perplexity, או Gemini '
        '"מי מספק [השירות שלך] ב[העיר שלך]?" — האם הם ימצאו אותך? '
        'זה בדיוק מה שמודד ומשפר הלשונית הזו.'
    )

    add_heading(doc, 'ציון נראות AI', level=3)
    add_bullet(doc, 'ציון 0-100 + דירוג A עד F')
    add_bullet(doc, 'שינוי מהביקורת הקודמת')

    add_heading(doc, 'תת-לשוניות', level=3)
    add_bullet(doc, 'בעיות ותיקונים: critical / warning / info — עם הסבר ופתרון ספציפי')
    add_bullet(doc, 'Schema Code: קוד JSON-LD מוכן להדבקה באתרך')
    add_bullet(doc, 'תוכן FAQ: שאלות ותשובות מוכנות להוספה לאתר')
    add_bullet(doc, 'תיאור עסקי: 120 מילה מותאמות לAI, מוכן להעתקה')
    add_bullet(doc, 'רשימת פעולות: צ\'קליסט ממוין — schema.org, Google Business, ועוד')

    # Tab 8: History
    add_heading(doc, 'לשונית 8: היסטוריה (History)', level=2)
    add_paragraph(doc,
        '12 החודשים האחרונים של הפעילות הפרסומית שלך, מסוכמים לפי חודש: '
        'הוצאה, ROAS, קונברסיות, CPA, CTR, קליקים, חשיפות, ציון GEO, '
        'ו-30 הפעולות הבולטות של Vigmis אותו חודש.'
    )

    # Tab 9: Decision Protocols
    add_heading(doc, 'לשונית 9: פרוטוקולי החלטה (Decision Protocols)', level=2)
    add_paragraph(doc,
        'Vigmis לא פועל בלי שתדע. כל החלטה משמעותית — '
        'שינוי תקציב, השהיית קמפיין, עדכון אסטרטגיה — '
        'מתועדת כאן עם הסבר מלא.'
    )
    add_bullet(doc, 'סוגי פרוטוקולים: אישור אסטרטגיה, שינוי תקציב, השהיית קמפיין, חידוש, הגדלה, רענון קריאטיבי, ניתוח קהל, התראת קיפאון')
    add_bullet(doc, 'סטטוס: ממתין → בדיון → מאושר / נדחה / פג תוקף (7 ימים)')
    add_bullet(doc, 'Thread שיחה — שאל שאלות, Vigmis ישיב')
    add_bullet(doc, 'אישור או דחייה — אתה שולט בכל החלטה')
    add_bullet(doc, 'payload מלא — מה בדיוק Vigmis תכנן לשנות')

    # Tab 10: Social
    add_heading(doc, 'לשונית 10: סושיאל (Social)', level=2)
    add_paragraph(doc,
        'Vigmis מנהל את הנוכחות האורגנית ברשתות החברתיות בנפרד מהפרסום בתשלום.'
    )

    add_heading(doc, 'יצירת תוכן שבועי', level=3)
    add_bullet(doc, 'Vigmis מייצר פוסטים לפייסבוק, אינסטגרם, ו-TikTok')
    add_bullet(doc, 'ברייף מותאם לעסק — הנחיות, נושאים, טון הקול')
    add_bullet(doc, 'תמונה AI לכל פוסט (פייסבוק/אינסטגרם)')
    add_bullet(doc, 'hashtag אוטומטי')
    add_bullet(doc, 'פורמט מותאם לכל פלטפורמה')
    add_bullet(doc, 'AI disclosure label (לפי חוקי EU AI Act)')

    add_heading(doc, 'מצבי אישור', level=3)
    add_table(doc,
        ['מצב', 'הסבר'],
        [
            ['אוטומטי', 'Vigmis מפרסם מיד — ללא אישורך'],
            ['לסקירה (ברירת מחדל)', 'Vigmis מחכה לאישורך לפני פרסום'],
            ['קפדני', 'כל פוסט עובר אישור, כולל עריכה חובה'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'ניהול תגובות', level=3)
    add_bullet(doc, 'Vigmis קורא את התגובות על הפוסטים שלך')
    add_bullet(doc, 'ממיין לפי דחיפות: דחוף / ערך גבוה / ניטרלי')
    add_bullet(doc, 'מציע תגובה מנוסחת מראש לכל תגובה')
    add_bullet(doc, 'זיהוי "משבר" — מתריע אם יש הצטברות של תגובות שליליות')
    add_bullet(doc, 'פעולות: שלח תגובה / התעלם / הסתר')

    add_heading(doc, 'אנליטיקת סושיאל', level=3)
    add_bullet(doc, 'לייקים, תגובות, שיתופים, reach, impressions, engagement rate')
    add_bullet(doc, 'סינון לפי טווח תאריכים')

    # Tab 11: Settings
    add_heading(doc, 'לשונית 11: הגדרות (Settings)', level=2)
    add_bullet(doc, 'פרטי עסק: שם, כתובת אתר, מטרה, אזור גיאוגרפי')
    add_bullet(doc, 'העדפות התראה: אימייל + WhatsApp')
    add_bullet(doc, 'סוגי התראות: חריגת הוצאה, נפילת CTR, עייפות קריאטיבית, מיצוי תקציב, שגיאת קמפיין')
    add_bullet(doc, 'אחוז שולי רווח (לחישוב ROAS אמיתי)')
    add_bullet(doc, 'חיבורי פלטפורמות: סטטוס + אפשרות התחברות מחדש')
    add_bullet(doc, 'Google Analytics 4: בחירת Property + סנכרון')
    add_bullet(doc, 'חברי צוות: הזמנה + ניהול הרשאות')
    add_bullet(doc, 'חיוב: שדרוג לScale + ניהול מנוי')
    add_bullet(doc, 'מחיקת חשבון')

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 4: Ask Vigmis
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 4: Ask Vigmis — הצ\'אט עם מנהל השיווק', level=1)

    add_paragraph(doc,
        'בכל עמוד בדשבורד תמצא כפתור "Ask Vigmis". '
        'זה לא צ\'אטבוט גנרי — זה מנהל שיווק שמכיר אותך אישית.'
    )

    add_heading(doc, 'מה Vigmis יודע עליך', level=2)
    add_bullet(doc, 'האסטרטגיה המלאה שנבנתה בonboarding')
    add_bullet(doc, 'ניתוח האתר שלך')
    add_bullet(doc, 'כל הקמפיינים הפעילים — גיל הקמפיין, תקופת למידה, ביצועים')
    add_bullet(doc, 'נתוני GA4 — מדדי KPI חיים')
    add_bullet(doc, 'פרוטוקולי החלטה פתוחים')
    add_bullet(doc, 'פעולות AI אחרונות')
    add_bullet(doc, 'פוסטים חברתיים')
    add_bullet(doc, 'התראות חדשות')

    add_heading(doc, 'מה Vigmis יכול לענות', level=2)
    add_bullet(doc, '"למה הביצועים איטיים?" — Vigmis יסביר שהקמפיין עדיין בתקופת למידה ומתי לצפות לשיפור')
    add_bullet(doc, '"כמה כסף להשקיע?" — ינתח לפי השוק שלך ויציע טווח')
    add_bullet(doc, '"מה לפרסם עכשיו?" — יבדוק מה עובד טוב ומה מומלץ לנסות')
    add_bullet(doc, '"למה ה-ROAS ירד?" — יבחן את הנתונים ויסביר')
    add_bullet(doc, 'שאלות אסטרטגיות — "כדאי להתרחב לTikTok?", "מתי לסיים ניסוי A/B?"')

    add_heading(doc, 'זיכרון שיחה', level=2)
    add_paragraph(doc,
        'Vigmis שומר 20 הודעות אחרונות בזיכרון — '
        'אתה לא צריך לחזור ולהסביר הכל מחדש בכל שיחה.'
    )

    add_info_box(doc,
        'Vigmis מחזיק עמדה ומסביר אותה. הוא לא יאמר "כן" לכל דבר — '
        'אם אתה רוצה לעשות משהו שלפי הנתונים לא ייעבד, הוא יאמר לך למה.'
    )

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 5: הדוח השבועי
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 5: הדוח השבועי', level=1)

    add_paragraph(doc,
        'פעם בשבוע, Vigmis שולח לך דוח אוטומטי שכתוב בשפה אנושית — '
        'לא טבלאות ונתונים יבשים, אלא נרטיב ברור:'
    )
    add_bullet(doc, 'מה קרה השבוע בפרסום שלך')
    add_bullet(doc, 'מה עבד טוב ומדוע')
    add_bullet(doc, 'מה Vigmis שינה ולמה')
    add_bullet(doc, 'מה לצפות לשבוע הבא')
    add_bullet(doc, 'המלצה אחת לפעולה שכדאי שתעשה')

    add_heading(doc, 'אמצעי הגעה', level=2)
    add_bullet(doc, 'אימייל (SendGrid)')
    add_bullet(doc, 'WhatsApp (לדוחות קצרים וחשובים)')

    add_heading(doc, 'התאמה אישית', level=2)
    add_bullet(doc, 'תדירות: יומי או שבועי (לפי העדפתך)')
    add_bullet(doc, 'שעת שליחה ו-timezone')
    add_bullet(doc, 'WhatsApp עבור לידים חשובים (תגובות בעלות פוטנציאל גבוה)')

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 6: קריאייטיב ותוכן
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 6: קריאייטיב ותוכן — הפרטים המלאים', level=1)

    add_heading(doc, 'Creative Director AI', level=2)
    add_paragraph(doc,
        'לפני שמייצרים כל קריאטיב — Vigmis מפעיל מנהל יצירתי AI '
        'שמנתח את הקשר המלא של העסק:'
    )
    add_bullet(doc, 'ניתוח האתר וחומרי המותג')
    add_bullet(doc, 'ביצועי הקריאטיבים הקודמים')
    add_bullet(doc, 'האסטרטגיה ויעד הקמפיין')
    add_bullet(doc, 'DNA המותג (צבעים, פונטים, טון)')

    add_paragraph(doc,
        'ואז בונה ברייף מפורט שמתאים לסוג הקריאטיב: '
        'לסרטון אווטאר — סקריפט שמתאים להגיית אנושית; '
        'לתמונה — הרכב ויזואלי ועיגון טקסט; '
        'לאנימציה — זרימת פריימים ומוסיקה.'
    )

    add_heading(doc, 'סרטוני אווטאר (HeyGen)', level=2)
    add_paragraph(doc,
        'Vigmis כותב סקריפט מלא ושולח ל-HeyGen שמייצר סרטון '
        'עם דמות AI שמדברת את הסקריפט בהתאמה מלאה. '
        'מתאים למודעות שמסבירות מוצר/שירות באופן אישי.'
    )

    add_heading(doc, 'סרטונים קולנועיים (Replicate)', level=2)
    add_paragraph(doc,
        'סרטוני וידאו קינמטיים ברזולוציית 1080p, מתאימים למודעות '
        'YouTube, Meta, ופרסום כללי ברשת. '
        'Vigmis מייצר ה-prompt ומעביר ל-Replicate שמייצר את הסרטון.'
    )

    add_heading(doc, 'סרטוני אנימציה (Replicate)', level=2)
    add_paragraph(doc,
        'סרטוני אנימציה ב-720p בפורמט אנכי (9:16), '
        'מותאמים לReels, TikTok, ו-Stories. '
        'מחיר נמוך יותר, הפקה מהירה יותר.'
    )

    add_heading(doc, 'ייצור תמונות (GPT-image-1)', level=2)
    add_paragraph(doc,
        'Vigmis מייצר שלוש תמונות מ-prompt מותאם אישית, '
        'AI Critic מדרג כל אחת (0-100) ובוחר הטובה ביותר. '
        'אם הציון נמוך מ-75, מנסה שוב אוטומטית.'
    )

    add_heading(doc, 'Brand DNA — זהות מותגית', level=2)
    add_bullet(doc, 'צבעי המותג (hex codes)')
    add_bullet(doc, 'פונטים מאושרים')
    add_bullet(doc, 'אלמנטים שלא לשנות (do_not_change)')
    add_bullet(doc, 'סגנונות מאושרים')
    add_bullet(doc, 'לוגו מוזרק אוטומטית לכל סרטון')

    add_heading(doc, 'לולאת משוב קריאטיבי', level=2)
    add_paragraph(doc,
        'כאשר קריאטיב מאושר ומפורסם, Vigmis מעקב אחרי הביצועים שלו. '
        'הנושאים, הוקים, וסגנונות שמביאים הכי הרבה קליקים והמרות — '
        'מוזרמים חזרה לCreative Director, כדי שהקריאטיב הבא יהיה טוב יותר.'
    )

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 7: מנגנוני אבטחה ובטיחות
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 7: מנגנוני אבטחה ובטיחות', level=1)

    add_heading(doc, 'מדיניות תוכן', level=2)
    add_paragraph(doc,
        'Vigmis מסרב ליצור תוכן בקטגוריות הבאות, ללא יוצא מן הכלל:'
    )
    add_bullet(doc, 'נשק וזיון')
    add_bullet(doc, 'סמים (לרבות קנאביס ללא היתר)')
    add_bullet(doc, 'תרופות ללא אישור רגולטורי')
    add_bullet(doc, 'מפעלי הימורים')
    add_bullet(doc, 'פירמידות / MLM')
    add_bullet(doc, 'הסתה לשנאה')

    add_info_box(doc,
        'סירוב אתי = הגנה על הפלטפורמה כולה ועל כל הלקוחות. '
        'Vigmis היא publisher, לא רק כלי — וכ-publisher היא אחראית לתוכן שהיא מפרסמת.'
    )

    add_heading(doc, 'כפתור עצירת חירום', level=2)
    add_paragraph(doc,
        'בכל רגע, מהמסך הראשי, ניתן ללחוץ על "עצירת חירום" — '
        'כל הקמפיינים נעצרים מיידית. '
        'WhatsApp alert נשלח אוטומטית לכל שינוי קריטי.'
    )

    add_heading(doc, 'שקיפות מלאה — Decision Protocols', level=2)
    add_paragraph(doc,
        'Vigmis מתעד כל החלטה: מה שונה, מדוע, מה התוצאה הצפויה. '
        'אתה תמיד יכול לראות מה Vigmis עשה ולא עשה, ומתי.'
    )

    add_heading(doc, 'בידוד נתונים (Tenant Isolation)', level=2)
    add_paragraph(doc,
        'כל עסק רואה אך ורק את הנתונים שלו. '
        'אין אפשרות גישה לנתוני עסק אחר — הבידוד מיושם ברמת מסד הנתונים.'
    )

    add_heading(doc, 'אימות ואבטחת חשבון', level=2)
    add_bullet(doc, 'Clerk JWT — אימות מאובטח בכל בקשת API')
    add_bullet(doc, 'סיסמאות Google OAuth — ללא אחסון סיסמה')
    add_bullet(doc, 'טוקני OAuth מוצפנים בזמן שמירה')
    add_bullet(doc, 'rate limiting — 10 ניסיונות כושלים נועלים את החשבון')
    add_bullet(doc, 'אימות HMAC לכל webhook (Stripe, Meta, Shopify)')

    add_heading(doc, 'Attestations — הצהרות משפטיות', level=2)
    add_paragraph(doc,
        'בכניסה ראשונה, הלקוח חותם על שלוש הצהרות מאובטחות:'
    )
    add_bullet(doc, 'דיוק המידע העסקי + אישור תנאי שימוש')
    add_bullet(doc, 'אישור תנאי שירות (ToS + AUP)')
    add_bullet(doc, 'הסכמת גילוי AI — אני יודע שהתוכן נוצר ע"י AI (לפי EU AI Act)')

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 8: מחירים ותוכניות
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 8: מחירים ותוכניות', level=1)

    add_heading(doc, 'תוכניות מנוי', level=2)
    add_table(doc,
        ['', 'Core (חינם)', 'Scale ($49/חודש)'],
        [
            ['ניהול קמפיינים', 'כן', 'כן'],
            ['אנליטיקה', 'כן', 'כן'],
            ['Ask Vigmis', 'כן', 'כן'],
            ['פוסטים חברתיים', 'כן', 'כן'],
            ['קריאייטיב (קרדיטים)', 'לא', '10 קרדיטים/חודש'],
            ['מספר מושבים (Seats)', '1', '3'],
            ['אופטימיזציה בלתי מוגבלת', 'לא', 'כן'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'קרדיטים לקריאייטיב (Scale)', level=2)
    add_table(doc,
        ['סוג קריאטיב', 'עלות', 'קרדיטים'],
        [
            ['סרטון אווטאר', '$15', '1 קרדיט'],
            ['סרטון קולנועי', '$12', '1 קרדיט'],
            ['סרטון אנימציה', '$8', '1 קרדיט'],
            ['תמונת AI', '$5', '1 קרדיט = 5 תמונות'],
        ]
    )
    doc.add_paragraph()

    add_bullet(doc, 'קרדיטים נצרכים בעת שליחת הבקשה')
    add_bullet(doc, 'קרדיטים מוחזרים אוטומטית אם הייצור נכשל')
    add_bullet(doc, 'אין העברת קרדיטים בין חודשים')
    add_bullet(doc, 'רוויזיות 1-2 חינם, 3-5 ב-50%, 6+ קריאטיב חדש')

    add_heading(doc, 'דמי ניהול', level=2)
    add_paragraph(doc,
        '7% מהוצאת הפרסום המנוהלת — נגבה חודשית. '
        'לדוגמה: תקציב חודשי $1,000 ← דמי ניהול $70. '
        'חשבונית חודשית + Stripe Customer Portal לניהול ולביטול.'
    )

    add_info_box(doc,
        'Vigmis מנהל רק את אחוז התקציב שאישרת — '
        'בונוס: תוכל להגדיר שVigmis ינהל רק 50% מהתקציב (שאר 50% תנהל בעצמך).'
    )

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 9: שירותים ותוכנות שמשתתפים ב-Vigmis
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 9: שירותים ותוכנות — כל מה שמשתתף ב-Vigmis', level=1)

    add_paragraph(doc,
        'Vigmis בנויה על שכבה של שירותים מקצועיים, כל אחד מומחה בתחומו. '
        'יחד הם יוצרים מנהל פרסום שלם.'
    )

    add_heading(doc, 'תשתית ומסד נתונים', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['Supabase', 'מסד נתונים PostgreSQL, אחסון קבצים, Row-Level Security', 'supabase.com'],
            ['Railway', 'אחסון שרת — API ו-Web App', 'railway.app'],
            ['Cloudflare R2', 'גיבויי נתונים', 'cloudflare.com'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'אימות ומשתמשים', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['Clerk', 'ניהול משתמשים, התחברות, Google OAuth, JWT', 'clerk.com'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'בינה מלאכותית', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['Anthropic Claude', 'מנוע AI ראשי — אסטרטגיה, צ\'אט, ניתוח, Creative Director, Decision Protocols', 'anthropic.com'],
            ['OpenAI gpt-image-1', 'יצירת תמונות פרסום, AI Critic', 'openai.com'],
            ['Perplexity Sonar Pro', 'מחקר שוק בזמן אמת — גישה לאינטרנט', 'perplexity.ai'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'ייצור קריאייטיב', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['HeyGen', 'יצירת סרטוני אווטאר — דמות AI שמדברת', 'heygen.com'],
            ['Replicate', 'יצירת סרטוני Cinematic ו-Animation', 'replicate.com'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'פרסום ואנליטיקה', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['Google Ads API', 'יצירה וניהול קמפיינים ב-Google', 'developers.google.com'],
            ['Meta Marketing API', 'יצירה וניהול קמפיינים ב-Facebook ו-Instagram', 'developers.facebook.com'],
            ['TikTok Ads API', 'ניהול קמפיינים ב-TikTok (בפיתוח)', 'business.tiktok.com'],
            ['Google Analytics 4', 'נתוני אנליטיקה מאתר הלקוח', 'analytics.google.com'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'מסחר ותשלומים', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['Shopify API', 'מעקב הזמנות, AOV, True ROAS', 'shopify.com'],
            ['Stripe', 'עיבוד תשלומים, מנוי Scale, דמי ניהול', 'stripe.com'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'תקשורת והתראות', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['WhatsApp Business API (Twilio)', 'התראות דחופות, דוחות שבועיים', 'business.whatsapp.com'],
            ['SendGrid', 'שליחת אימיילים — דוחות, הזמנות, התראות', 'sendgrid.com'],
            ['NewsAPI', 'סריקת חדשות שוק שבועית', 'newsapi.org'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'ניטור וניתוח', level=2)
    add_table(doc,
        ['שירות', 'תפקיד', 'כתובת'],
        [
            ['PostHog', 'ניתוח התנהגות משתמשים בתוך Vigmis', 'posthog.com'],
            ['Instatus', 'ניטור זמינות המערכת', 'instatus.com'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'סטק טכנולוגי (מידע למתעניינים)', level=2)
    add_table(doc,
        ['שכבה', 'טכנולוגיה'],
        [
            ['ממשק משתמש', 'Next.js App Router, TypeScript'],
            ['שרת API', 'Fastify, Node.js, TypeScript'],
            ['עיצוב', 'Tailwind CSS'],
            ['לוקליזציה', 'next-intl — 13 שפות (עברית, ערבית, אנגלית ועוד)'],
            ['RTL', 'תמיכה מלאה בעברית וערבית'],
        ]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 10: מנועי הבינה — Intelligence Engines
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 10: מנועי הבינה — Intelligence Engines', level=1)

    add_paragraph(doc,
        'ברוב כלי הפרסום — ה-AI מקבל החלטה, מיישם אותה, ומעולם לא בודק אם היא עבדה. '
        'Vigmis בנה שכבת בינה שסוגרת את הלולאה: כל החלטה נמדדת 10 ימים אחריה, '
        'וכל תוצאה משנה את ההחלטה הבאה. '
        'לאורך זמן, Vigmis יודע אילו החלטות עובדות עבור העסק הספציפי שלך.'
    )

    add_info_box(doc,
        'מנועי הבינה מופעלים בהדרגה לפי רמת ה-Data Maturity שלך. '
        'עסק חדש מקבל מנועים בסיסיים; עסק עם 3+ חודשי נתונים מקבל את כל הארסנל.'
    )

    # ── Data Maturity Score ──
    add_heading(doc, 'Data Maturity Score — בשלות המידע', level=2)
    add_paragraph(doc,
        'לפני שVigmis מציג המלצה — הוא בודק כמה נתונים יש לו לעמוד עליהם. '
        'ציון הבשלות (1-5) קובע אילו מנועים מופעלים:'
    )
    add_table(doc,
        ['רמה', 'תנאי', 'מה מופעל'],
        [
            ['1 — מתחיל', '< 14 ימים או < 30 קליקים', 'אופטימיזציה בסיסית בלבד'],
            ['2 — לומד', '14-30 ימים + 30-100 קליקים', '+ ניתוח A/B, גילוי קהלים'],
            ['3 — נתונים', '> 30 ימים + 100 קליקים + GA4', '+ Strategic Brain שבועי, Explore/Exploit'],
            ['4 — בשל', '> 90 ימים + 500 קליקים + 2 פלטפורמות', '+ Portfolio Allocator, Three-ROAS'],
            ['5 — מומחה', '> 180 ימים + 2,000 קליקים + Shopify', '+ כל המנועים, True Incrementality'],
        ]
    )
    doc.add_paragraph()
    add_paragraph(doc,
        'הציון מחושב מחדש כל שבוע (יום שני 10:00 UTC). '
        'Vigmis מדווח ב-Ask Vigmis מה רמתך הנוכחית ומה חסר כדי לעלות שלב.'
    )

    # ── Strategic Brain ──
    add_heading(doc, 'Strategic Brain — ניתוח שבועי עמוק', level=2)
    add_paragraph(doc,
        'פעם בשבוע (יום שני 09:00 UTC) Vigmis מריץ ניתוח אסטרטגי מלא. '
        'בניגוד לאופטימיזציה היומית שמסתכלת על מטריקות, '
        'ה-Strategic Brain שואל: "האסטרטגיה הכוללת עדיין נכונה?"'
    )
    add_heading(doc, 'מה ה-Strategic Brain מנתח', level=3)
    add_bullet(doc, 'CTR ו-ROAS אמיתיים מ-7 הימים האחרונים (מתוך snapshots ב-audit log — לא דיווח פלטפורמה)')
    add_bullet(doc, 'נתוני GA4 כ-"אמת" — אטריביוציה עצמאית מול דיווח Meta/Google')
    add_bullet(doc, 'Three-ROAS: platform_reported / ga4_attributed / incremental_estimate (ברמה 4+)')
    add_bullet(doc, 'Regime Detection: השוואת CTR 7 ימים מול 60 ימים — אם יחס < 75%, מסמן "degrading"')
    add_bullet(doc, 'תוצאות החלטות קודמות — אילו decision protocols עבדו ואילו לא')
    add_bullet(doc, 'ממוצע batting average לפי סוג החלטה (scale up, pause, budget shift, creative refresh)')
    add_bullet(doc, 'השערות פתוחות (hypotheses) שעוד לא נבדקו')

    add_heading(doc, 'תוצרי ה-Strategic Brain', level=3)
    add_bullet(doc, 'פסיקת תיק (portfolio verdict): on_track / at_risk / scaling / degrading')
    add_bullet(doc, 'regime_signal: stable / improving / degrading')
    add_bullet(doc, 'key_insights: 3-5 תובנות בשפה אנושית')
    add_bullet(doc, 'recommended_focus: מה לעשות השבוע הבא')
    add_bullet(doc, 'new_hypotheses: השערות חדשות לבדיקה (עד 10 פתוחות בו-זמנית, FIFO)')

    # ── Outcome Tracker ──
    add_heading(doc, 'Outcome Tracker — לולאת המשוב', level=2)
    add_paragraph(doc,
        'כל פרוטוקול החלטה שאושר (שינוי תקציב, הגדלת קמפיין, השהיה, חידוש) '
        'מסומן לבדיקה 10 ימים אחרי האישור. '
        'Vigmis בודק: "האם הביצועים השתפרו, הידרדרו, או לא השתנו?"'
    )
    add_table(doc,
        ['שלב', 'מה קורה'],
        [
            ['אישור פרוטוקול', 'Vigmis מסמן check_after = now + 10 days'],
            ['יום 10', 'outcome-tracker משווה CTR/ROAS 7 ימים לפני ו-7 ימים אחרי'],
            ['תוצאה', 'improved (>5%) / worsened (<-5%) / neutral / insufficient_data'],
            ['עדכון batting average', 'כל סוג החלטה צובר ממוצע הצלחה עדכני'],
            ['זרימה לפנים', 'Strategic Brain קורא את התוצאות ומשנה המלצות בהתאם'],
        ]
    )
    doc.add_paragraph()
    add_paragraph(doc,
        'Vigmis דורש לפחות 3 snapshots לפני ו-3 snapshots אחרי ההחלטה לפני שהוא מסיק מסקנה. '
        'אם אין מספיק נתונים — הוא מסמן insufficient_data ולא מסיק מסקנות שגויות.'
    )
    add_bold_paragraph(doc, 'שקיפות: ',
        'כל תוצאת מדידה מופיעה בדוח השבועי תחת "החלטות שנמדדו השבוע".')

    # ── Decision Quality / Batting Average ──
    add_heading(doc, 'Decision Quality — ממוצע הצלחת החלטות', level=2)
    add_paragraph(doc,
        'Vigmis מחשב batting average לכל סוג החלטה:'
    )
    add_bullet(doc, 'campaign_scale — כמה פעמים הגדלת תקציב עבדה (CTR/ROAS עלה)')
    add_bullet(doc, 'campaign_pause — כמה פעמים עצירת קמפיין הייתה נכונה')
    add_bullet(doc, 'budget_change — ממוצע שינוי תקציב')
    add_bullet(doc, 'portfolio_reallocation — ממוצע העברה בין פלטפורמות')
    add_paragraph(doc,
        'הנתונים מוזנים ל-Ask Vigmis ולדוח השבועי, כדי ש-Vigmis יוכל לכמת: '
        '"החלטות scale up שלנו מצליחות ב-72% מהמקרים — מעל ממוצע ה-AI industry."'
    )

    # ── Portfolio Allocator ──
    add_heading(doc, 'Portfolio Allocator — הקצאת הון חכמה', level=2)
    add_paragraph(doc,
        'רוב כלי הפרסום מראים לך ROAS כפי שפלטפורמת הפרסום מדווחת. '
        'הבעיה: כל פלטפורמה מייחסת לעצמה יותר ממה שהיא הרוויחה. '
        'Portfolio Allocator מסתמך על Google Analytics 4 — מקור אמת חיצוני לשתי הפלטפורמות.'
    )
    add_heading(doc, 'איך זה עובד', level=3)
    add_bullet(doc, 'Vigmis שולף נתוני ROAS מ-GA4 לפי medium/source (cpc=Google, paid_social=Meta)')
    add_bullet(doc, 'אם הפער בין הפלטפורמות גדול מ-80% — Vigmis יוצר פרוטוקול הקצאה מחדש')
    add_bullet(doc, 'ממוצע הקצאה מחדש לא יעלה על 30% מהתקציב בבת אחת (safe cap)')
    add_bullet(doc, 'מינימום שינוי: $5/יום (אחרת לא שווה פרוטוקול)')
    add_bullet(doc, 'אידמפוטנטי: לא ייצור פרוטוקול חדש אם כבר יש אחד ממתין, או אם אושר תוך 30 ימים')
    add_bullet(doc, 'דורש Data Maturity 4+ (90+ ימים, 2 פלטפורמות, GA4 מחובר)')
    add_info_box(doc,
        'Three-ROAS display: בכל פרוטוקול scale/reallocation — Vigmis מציג שלושה מספרים: '
        'מה Meta/Google מדווחים, מה GA4 מיחס, ומה Vigmis מעריך כ-incremental (תוספת אמיתית). '
        'רוב כלי הפרסום מראים רק את המספר הראשון.'
    )

    # ── Hypothesis Engine ──
    add_heading(doc, 'Hypothesis Engine — מנגנון ההשערות', level=2)
    add_paragraph(doc,
        'הבעיה הגדולה של A/B testing: מה בודקים? '
        'Vigmis פותר אותה ע"י השערות שנוצרות מהנתונים ועוברות לקריאייטיב:'
    )
    add_bullet(doc, 'Strategic Brain מנתח ביצועי קמפיינים, תגובות תגובות, ומגמות שוק')
    add_bullet(doc, 'הוא מייצר השערות כמו: "קהל 35-44 מגיב טוב יותר להוק כלכלי מאשר רגשי"')
    add_bullet(doc, 'ההשערות מאוחסנות ב-client_settings.hypotheses (מקסימום 10 פתוחות, FIFO)')
    add_bullet(doc, 'Creative Director AI קורא את ההשערות ומנסה ליישם אותן בקריאייטיב הבא')
    add_bullet(doc, 'אחרי בדיקה (A/B test, outcome measurement) — ההשערה נסגרת עם תוצאה')
    add_paragraph(doc,
        'דוגמה לשרשרת: Strategic Brain כותב השערה → Creative Director כותב סרטון שבודק אותה → '
        'A/B test מראה שהיא עבדה → Strategic Brain מסמן confirmed ומייצר השערה הבאה.'
    )

    # ── Explore/Exploit ──
    add_heading(doc, 'Explore/Exploit — מתי לגדול, מתי לבדוק', level=2)
    add_paragraph(doc,
        'בעיה קלאסית בפרסום: אחרי הגדלת תקציב שעבדה, יש פיתוי להגדיל עוד ועוד. '
        'אבל לפעמים צריך לבדוק קריאטיב חדש — לא רק לנצל מה שעובד.'
    )
    add_bullet(doc, 'אם אושר scale_up תוך 30 ימים + אין A/B test פעיל → Vigmis מציע לפתוח A/B test')
    add_bullet(doc, 'הפרוטוקול מסביר: "ביצועים חזקים עכשיו = זמן אידיאלי לבדיקה עם control חזק"')
    add_bullet(doc, 'זה protect mode: לא ממליץ לשנות את מה שעובד, רק לבדוק לצדו')
    add_bullet(doc, 'ניתן לכבות: ENABLE_EXPLORE_EXPLOIT=false')

    # ── Regime Detection ──
    add_heading(doc, 'Regime Detection — זיהוי שינוי מגמה', level=2)
    add_paragraph(doc,
        'CTR יורד יום אחד — זה רעש. CTR יורד 3 שבועות ברציפות — זה regime שינוי. '
        'Vigmis מבחין בין השניים:'
    )
    add_bullet(doc, 'משווה CTR ממוצע 7 ימים מול CTR ממוצע 60 ימים')
    add_bullet(doc, 'אם יחס < 75% → regime = "degrading", Vigmis מסמן בדוח השבועי')
    add_bullet(doc, 'Regime degrading → Strategic Brain ממליץ creative refresh ולא budget scale')
    add_bullet(doc, 'Regime stable/improving → Budget allocator רשאי להמליץ scale')

    # Feature flags summary
    add_heading(doc, 'בקרת מנועים — Feature Flags', level=2)
    add_table(doc,
        ['Flag', 'מנוע', 'ברירת מחדל'],
        [
            ['ENABLE_OUTCOME_TRACKER', 'לולאת המשוב + batting average', 'on'],
            ['ENABLE_PORTFOLIO_ALLOCATOR', 'Portfolio Allocator (GA4 ROAS)', 'on'],
            ['ENABLE_HYPOTHESES', 'Hypothesis Engine', 'on'],
            ['ENABLE_EXPLORE_EXPLOIT', 'Explore/Exploit suggestion', 'on'],
        ]
    )
    doc.add_paragraph()
    add_info_box(doc,
        'כל מנוע ניתן לכיבוי בנפרד דרך Railway env vars. '
        'כיבוי מנוע לא מוחק נתונים היסטוריים — רק מפסיק הרצה חדשה.'
    )

    # ══════════════════════════════════════════════════════════════════════════════
    # פרק 11: מה בדרך
    # ══════════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_heading(doc, 'פרק 11: מה בדרך — בפיתוח', level=1)

    add_paragraph(doc,
        'Vigmis מתפתחת כל הזמן. הפיצ\'רים הבאים הם בשלבי פיתוח '
        'ויגיעו בחודשים הקרובים:'
    )

    add_heading(doc, 'TikTok — אינטגרציה מלאה', level=2)
    add_paragraph(doc,
        'הקוד מוכן ומחובר. ממתינים לאישור App Review של TikTok '
        '(צפי: קיץ 2026). '
        'ברגע שהאישור יתקבל — ניהול קמפיינים ב-TikTok יהיה זמין מיידית.'
    )

    add_heading(doc, 'חשבונות צוות (Multi-user)', level=2)
    add_paragraph(doc,
        'קוד ניהול הצוות מוכן (Scale: 3 מושבים, הזמנות במייל). '
        'בודקים E2E עם שני חשבונות בו-זמנית לפני השקה.'
    )

    add_heading(doc, 'Weekly Briefing — מצגת שבועית', level=2)
    add_paragraph(doc,
        'מצגת ויזואלית שבועית (slide deck) שמסכמת את השבוע בצורה ויזואלית. '
        'בנוסף לדוח הטקסטואלי הקיים.'
    )

    add_heading(doc, 'שיפורי Onboarding', level=2)
    add_paragraph(doc,
        'על בסיס נתוני PostHog על הפאנל, נשפר את חווית הכניסה הראשונה '
        'עבור משתמשים לא טכניים: '
        'overlay ברוכים הבאים, הסברים מוצגים יותר, ופחות עצירות.'
    )

    # ── Final Page: Summary ───────────────────────────────────────────────────────
    add_page_break(doc)
    add_heading(doc, 'סיכום: Vigmis בקליפת אגוז', level=1)

    add_paragraph(doc,
        'Vigmis היא לא "עוד אפליקציה". היא מחליפה מנהל פרסום שלם. '
        'היא מחוברת לפלטפורמות הפרסום שלך, מייצרת תוכן, '
        'מקבלת החלטות, מדווחת — ומסבירה הכל.'
    )

    add_table(doc,
        ['מה Vigmis עושה', 'איך'],
        [
            ['מנהלת קמפיינים ב-Google, Meta ו-TikTok', 'API מחובר ישיר, בזמן אמת'],
            ['בונה אסטרטגיה שיווקית', 'AI + מחקר שוק חי (Perplexity)'],
            ['מייצרת סרטונים ותמונות', 'HeyGen + Replicate + GPT-image-1'],
            ['כותבת מודעות ופוסטים', 'Claude AI + brand DNA + platform rules'],
            ['מנתחת ומייעלת', 'Daily optimization engine'],
            ['מתריעה בזמן אמת', 'WhatsApp + Email alerts'],
            ['מדווחת בשפה אנושית', 'AI narrative weekly briefing'],
            ['מגנה על המותג שלך', 'Content policy + attestations + audit trail'],
        ]
    )
    doc.add_paragraph()

    p_final = doc.add_paragraph()
    run = p_final.add_run(
        'vigmis.com'
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    doc = create_document()
    doc.save(OUTPUT_PATH)
    size = os.path.getsize(OUTPUT_PATH)
    print(f"SUCCESS: {OUTPUT_PATH}")
    print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
