"""
Generate professional Word documents for Vigmis legal files.
Run: python generate_word_docs.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ─────────────────────────────────────────────────────────────────

HE_FONT = 'Arial'       # supports Hebrew glyphs
EN_FONT = 'Calibri'


# ── Low-level RTL helpers ─────────────────────────────────────────────────────

def _set_para_rtl_props(paragraph, align_right=True):
    """Add w:bidi (and optionally w:jc right) to paragraph properties."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove duplicates
    for tag in ['w:bidi', 'w:jc']:
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    if align_right:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)


def _set_run_rtl_props(run, font_name=HE_FONT):
    """Add w:rtl and Hebrew language tag to a run."""
    rPr = run._r.get_or_add_rPr()
    # Remove duplicates
    for tag in ['w:rtl', 'w:lang']:
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:bidi'), 'he-IL')
    rPr.append(lang)
    if font_name:
        run.font.name = font_name


def make_rtl(paragraph, font_name=HE_FONT):
    """Make a paragraph (and all its existing runs) RTL."""
    _set_para_rtl_props(paragraph)
    for run in paragraph.runs:
        _set_run_rtl_props(run, font_name)


# ── Document-level helpers ────────────────────────────────────────────────────

def set_doc_rtl_defaults(doc):
    """Set Normal style defaults to RTL so new paragraphs inherit."""
    normal = doc.styles['Normal']
    # Paragraph defaults
    pPr = normal.element.get_or_add_pPr()
    for tag in ['w:bidi', 'w:jc']:
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)
    # Run defaults
    rPr = normal.element.get_or_add_rPr()
    for tag in ['w:rtl', 'w:lang']:
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:bidi'), 'he-IL')
    rPr.append(lang)


# ── High-level paragraph helpers (Hebrew doc) ─────────────────────────────────

def he_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    make_rtl(p)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def he_para(doc, text, bold=False, italic=False, size=11, space_after=6, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    make_rtl(p)
    return p


def he_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    make_rtl(p)
    return p


def he_table_row(table, cells, bold_first=False):
    """Add a row to a table; all cells get RTL."""
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        # Clear and set text through a run (so we can control RTL on the run)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        run.font.name = HE_FONT
        if bold_first and i == 0:
            run.bold = True
        make_rtl(cell.paragraphs[0])


def he_table_header(table, labels):
    """Style the first (header) row of a table."""
    row = table.rows[0]
    for cell, text in zip(row.cells, labels):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = HE_FONT
        make_rtl(cell.paragraphs[0])


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '94A3B8')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — Hebrew Executive Summary
# ══════════════════════════════════════════════════════════════════════════════

def build_hebrew_doc():
    doc = Document()
    set_doc_rtl_defaults(doc)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Title ──
    title = doc.add_heading('VIGMIS — מבנה משפטי ותאגידי', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_rtl(title)

    sub = doc.add_paragraph('סיכום מנהלים — גרסה מלאה')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].italic = True
    make_rtl(sub)

    meta = doc.add_paragraph('תאריך עריכה: יוני 2026   |   מוכן על ידי: עמיחי קרופיק, מייסד   |   סטטוס: טיוטת עבודה')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    make_rtl(meta)

    add_hr(doc)

    # ── 1. רקע ──
    he_heading(doc, '1. רקע ומטרת המסמך', level=1)
    he_para(doc,
        'מסמך זה מתאר את המבנה התאגידי שנבחר עבור פעילות VIGMIS, מנמק את ההחלטות שהתקבלו, '
        'ומשמש כבסיס לעבודה מול יועצים משפטיים, רואי חשבון ורשויות מס בישראל ובארצות הברית.')
    he_para(doc,
        'VIGMIS היא פלטפורמת SaaS לניהול שיווק דיגיטלי מבוסס AI, המיועדת לעסקים ברחבי העולם. '
        'הפלטפורמה מציעה יצירת תוכן שיווקי, ניהול קמפיינים, ניתוח ביצועים, ותקשורת עם לקוחות.')

    # ── 2. הצדדים ──
    he_heading(doc, '2. הצדדים', level=1)

    he_heading(doc, '2.1 טאורוס מנג\'מנט ואינווסטמנטס בע"מ', level=2)
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = 'Table Grid'
    he_table_header(t1, ['שדה', 'פרטים'])
    for r in [
        ('שם', 'Taurus Management and Investments Ltd.'),
        ('מספר חברה', '514565118'),
        ('מדינה', 'ישראל'),
        ('כתובת', 'מובשוביץ בנימין 25, הרצליה 4640525'),
        ('בעלות', '100% — עמיחי קרופיק'),
        ('תפקיד', 'חברת אם, בעלת IP'),
    ]:
        he_table_row(t1, r, bold_first=True)
    doc.add_paragraph()

    he_heading(doc, '2.2 VIGMIS US LLC', level=2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    he_table_header(t2, ['שדה', 'פרטים'])
    for r in [
        ('שם', 'VIGMIS US LLC'),
        ('מדינה', 'ארצות הברית — Wyoming'),
        ('בעלות', '100% — Taurus Management and Investments Ltd.'),
        ('UBO', 'עמיחי קרופיק'),
        ('תפקיד', 'חברת מכירות, גבייה ושיווק'),
    ]:
        he_table_row(t2, r, bold_first=True)
    doc.add_paragraph()

    # ── 3. מבנה הבעלות ──
    he_heading(doc, '3. מבנה הבעלות', level=1)
    tree = doc.add_paragraph()
    _set_para_rtl_props(tree, align_right=False)
    tree.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tree.add_run(
        'עמיחי קרופיק (פרט)\n'
        '          |\n'
        '          | 100%\n'
        '          v\n'
        'Taurus Management and Investments Ltd.\n'
        '(Israel — Company No. 514565118)\n'
        '          |\n'
        '          | 100%\n'
        '          v\n'
        'VIGMIS US LLC\n'
        '(Wyoming, USA)'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    _set_run_rtl_props(run, font_name=None)

    # ── 4. למה ──
    he_heading(doc, '4. למה נבחר מבנה זה?', level=1)

    he_heading(doc, '4.1 הצורך בישות אמריקאית', level=2)
    he_para(doc, 'פעילות VIGMIS מחייבת חשבון Stripe וחשבון בנק אמריקאי:')
    for item in [
        'Stripe — מעניק שירות אמין ומלא לחברות אמריקאיות',
        'Google Ads ו-Meta Ads — גבייה ותשלום קל יותר מחשבון בנק אמריקאי',
        'לקוחות גלובליים — חוזים מול ישות אמריקאית עדיפים',
        'שירותי ענן — AWS, Railway, OpenAI — גבייה קלה יותר מישות אמריקאית',
    ]:
        he_bullet(doc, item)

    he_heading(doc, '4.2 למה Wyoming?', level=2)
    for item in [
        'אין מס הכנסה מדינתי — 0% state income tax',
        'עלות הקמה ותפעול נמוכה — כ-$100-150 לשנה',
        'פרטיות גבוהה — אין חובת גילוי שמות בעלים',
        'חוקי LLC מודרניים ובהירים',
        'Registered Agent זמין וזול',
    ]:
        he_bullet(doc, item)

    he_heading(doc, '4.3 IP נשאר בישראל — למה?', level=2)
    for item in [
        'פיתוח מתבצע בישראל — הכרה עתידית בהכנסות R&D',
        'פשטות — העברת IP לארה"ב מחייבת הערכת שווי ומיסוי',
        'הגנה — IP בישות האם מוגן מתביעות נגד ה-LLC',
        'עקביות — כל הצוות והטכנולוגיה בישראל',
    ]:
        he_bullet(doc, item)

    # ── 5. חלוקת אחריות ──
    he_heading(doc, '5. חלוקת תחומי אחריות', level=1)

    he_heading(doc, 'טאורוס ישראל — אחראית על:', level=2)
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = 'Table Grid'
    he_table_header(t3, ['תחום', 'פירוט'])
    for r in [
        ('פיתוח מוצר', 'כל הקוד, אדריכלות המערכת, CI/CD'),
        ('R&D', 'מחקר ופיתוח AI, מודלים, Prompts'),
        ('Infrastructure', 'Railway, Supabase, Cloudflare, R2'),
        ('AI Providers', 'OpenAI, Anthropic, OpenRouter'),
        ('IP', 'קוד מקור, מותג, דומיינים, לוגו, AI Workflows'),
    ]:
        he_table_row(t3, r, bold_first=True)
    doc.add_paragraph()

    he_heading(doc, 'VIGMIS US LLC — אחראית על:', level=2)
    t4 = doc.add_table(rows=1, cols=2)
    t4.style = 'Table Grid'
    he_table_header(t4, ['תחום', 'פירוט'])
    for r in [
        ('גבייה', 'Stripe — ניהול מנויים, חיובים, Webhooks'),
        ('בנקאות', 'Mercury / Relay — חשבון בנק עסקי אמריקאי'),
        ('מכירות', 'Onboarding לקוחות, Trials, Conversions'),
        ('שיווק', 'Google Ads, Meta Ads, TikTok Ads'),
        ('חוזים', 'הלקוחות חותמים חוזה מול VIGMIS US LLC'),
    ]:
        he_table_row(t4, r, bold_first=True)
    doc.add_paragraph()

    doc.add_page_break()

    # ── 6. מודל הכנסות ──
    he_heading(doc, '6. מודל הכנסות וחלוקה', level=1)

    he_heading(doc, '6.1 הגדרות', level=2)
    p_def = doc.add_paragraph()
    p_def.paragraph_format.space_after = Pt(6)
    r1 = p_def.add_run('Gross Revenue')
    r1.bold = True
    r1.font.name = EN_FONT
    r2 = p_def.add_run(': סך החיובים ב-Stripe לפני כל ניכוי.')
    r2.font.name = HE_FONT
    make_rtl(p_def)

    p_net = doc.add_paragraph()
    p_net.paragraph_format.space_after = Pt(6)
    r3 = p_net.add_run('Net Revenue')
    r3.bold = True
    r3.font.name = EN_FONT
    r4 = p_net.add_run(' = Gross Revenue - Refunds - Chargebacks - Stripe Processing Fees (כ-2.9% + $0.30)')
    r4.font.name = HE_FONT
    make_rtl(p_net)

    he_heading(doc, '6.2 חלוקת Net Revenue', level=2)
    t5 = doc.add_table(rows=1, cols=3)
    t5.style = 'Table Grid'
    he_table_header(t5, ['צד', 'אחוז', 'הצדקה'])
    for r in [
        ('VIGMIS US LLC', '25%', 'כיסוי עלויות שיווק, מכירות, CRM, תפעול'),
        ('Taurus Management', '75%', 'תמלוג IP + שירותי פיתוח + תשתית'),
    ]:
        he_table_row(t5, r)
    doc.add_paragraph()

    # ── 7. שיקולי מס ──
    he_heading(doc, '7. שיקולי מס', level=1)

    he_heading(doc, '7.1 ארצות הברית', level=2)
    for item in [
        'Form 5472: חובה ל-Foreign-Owned SMLLC. קנס $25,000 לכל טרנזקציה שלא דווחה. לתאם עם CPA אמריקאי לפני הגשה ראשונה.',
        'Form 1120: VIGMIS US LLC מגישה דוח מס שנתי פדרלי.',
        'Wyoming: אין מס הכנסה מדינתי.',
    ]:
        he_bullet(doc, item)

    he_heading(doc, '7.2 ישראל', level=2)
    for item in [
        'CFC: כללי CFC של רשות המסים עשויים לחייב דיווח על הכנסות VIGMIS US בישראל. לבדוק עם יועץ מס ישראלי.',
        'העברות בין-חברתיות: 75% שמועבר לטאורוס — ניכוי מס במקור כפוף לאמנת המס ישראל-ארה"ב.',
        'Transfer Pricing: ה-75/25 חייב להיות מתועד כ-Arm\'s Length. ראה הסכם בין-חברתי.',
    ]:
        he_bullet(doc, item)

    # ── 8. משימות ──
    he_heading(doc, '8. רשימת משימות פתוחות', level=1)

    he_heading(doc, 'משפטי', level=2)
    for item in [
        'LLC Formation — הגשת Articles of Organization ב-Wyoming',
        'Operating Agreement — עם עורך דין Wyoming',
        'Intercompany Agreement — חתימה (ראה מסמך נפרד)',
        'Transfer Pricing Memo — עם הכנסות משמעותיות',
    ]:
        he_bullet(doc, item)

    he_heading(doc, 'מס ורגולציה', level=2)
    for item in [
        'CPA אמריקאי — Form 5472, Form 1120',
        'יועץ מס ישראלי — CFC, אמנת מס',
        'EIN — מה-IRS לאחר הקמת ה-LLC',
    ]:
        he_bullet(doc, item)

    he_heading(doc, 'בנקאות ותשלומים', level=2)
    for item in [
        'Mercury / Relay — פתיחת חשבון בנק עסקי',
        'Stripe — Live mode תחת VIGMIS US LLC',
    ]:
        he_bullet(doc, item)

    add_hr(doc)
    note = doc.add_paragraph(
        'מסמך זה נועד לצרכים עסקיים פנימיים ולשמש כבסיס לייעוץ מקצועי. '
        'אינו מהווה ייעוץ משפטי או מיסויי.'
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    make_rtl(note)

    doc.save(r'C:\vigmis\vigmis-main\LEGAL\מבנה_משפטי_ותאגידי.docx')
    print('Hebrew document saved')


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 — Intercompany Agreement (English) — unchanged
# ══════════════════════════════════════════════════════════════════════════════

def add_paragraph_en(doc, text, rtl=False, bold=False, italic=False, size=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_table_row_en(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                if bold_first and i == 0:
                    run.bold = True
                run.font.size = Pt(10)


def add_hr_en(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '94A3B8')
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_agreement_doc():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Title ──
    title = doc.add_heading('INTERCOMPANY LICENSE, SERVICES AND\nDISTRIBUTION AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    eff_date = doc.add_paragraph('Effective Date: ___________________________')
    eff_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eff_date.runs[0].font.size = Pt(12)
    eff_date.runs[0].italic = True

    add_hr_en(doc)

    # ── Parties ──
    doc.add_heading('PARTIES', level=1)
    pp = doc.add_paragraph('This Intercompany License, Services and Distribution Agreement (this ')
    pp.add_run('"Agreement"').bold = True
    pp.add_run(') is entered into as of the Effective Date by and between:')

    doc.add_paragraph()
    pt = doc.add_paragraph()
    pt.add_run('TAURUS MANAGEMENT AND INVESTMENTS LTD.').bold = True
    pt.add_run(', an Israeli private company, registered number 514565118, with its principal place of '
               'business at Benjamin Movshovitz 25, Herzliya 4640525, Israel (hereinafter ')
    pt.add_run('"Taurus"').bold = True
    pt.add_run(' or ')
    pt.add_run('"Licensor"').bold = True
    pt.add_run(');')

    doc.add_paragraph('and')

    pv = doc.add_paragraph()
    pv.add_run('VIGMIS US LLC').bold = True
    pv.add_run(', a Wyoming limited liability company, wholly owned by Taurus, with its registered agent '
               'in the State of Wyoming (hereinafter ')
    pv.add_run('"VIGMIS US"').bold = True
    pv.add_run(' or ')
    pv.add_run('"Licensee"').bold = True
    pv.add_run(').')

    doc.add_paragraph()
    pe = doc.add_paragraph()
    pe.add_run('Taurus').italic = True
    pe.add_run(' and ')
    pe.add_run('VIGMIS US').italic = True
    pe.add_run(' are each referred to herein individually as a ')
    pe.add_run('"Party"').bold = True
    pe.add_run(' and collectively as the ')
    pe.add_run('"Parties"').bold = True
    pe.add_run('.')

    add_hr_en(doc)

    # ── Recitals ──
    doc.add_heading('RECITALS', level=1)
    for r in [
        'WHEREAS, Taurus owns all right, title, and interest in the VIGMIS platform, including its software, '
        'artificial intelligence systems, brand, trademarks, domain names, and related intellectual property '
        '(collectively, the "IP");',
        'WHEREAS, Taurus desires to grant VIGMIS US a license to use the IP for the purpose of marketing, '
        'distributing, and selling access to the VIGMIS platform to end customers;',
        'WHEREAS, Taurus will continue to provide development, maintenance, and infrastructure services '
        'necessary to operate the platform;',
        'WHEREAS, the Parties desire to establish clear terms governing their intercompany relationship, '
        'including the allocation of revenues and costs, in a manner consistent with the arm\'s-length '
        'standard required under applicable tax law;',
    ]:
        p = doc.add_paragraph(r, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    pn = doc.add_paragraph()
    pn.add_run('NOW, THEREFORE').bold = True
    pn.add_run(', in consideration of the mutual covenants and agreements set forth herein, '
               'the Parties agree as follows:')

    add_hr_en(doc)

    # ── Helpers ──
    def art(t):
        h = doc.add_heading(t, level=1)
        h.paragraph_format.space_before = Pt(12)
        return h

    def sec(num, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(num + '  ').bold = True
        p.add_run(text)
        return p

    def sub(letter, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run('(' + letter + ')  ' + text)
        return p

    def defn(num, term, definition):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(num + '  ').bold = True
        p.add_run(term).bold = True
        p.add_run('  ' + definition)

    # ── Article 1 — Definitions ──
    art('ARTICLE 1 — DEFINITIONS')
    defn('1.1', '"Platform"',
         'means the VIGMIS SaaS platform, including all software, AI systems, APIs, user interfaces, '
         'mobile applications, backend services, and associated documentation developed and owned by Taurus.')
    defn('1.2', '"End Customer"',
         'means any third-party business or individual that subscribes to or purchases access to the '
         'Platform through VIGMIS US.')
    defn('1.3', '"Gross Revenue"',
         'means all amounts collected by VIGMIS US from End Customers in connection with their use of '
         'the Platform, including subscription fees, usage fees, and any other charges, before any deductions.')
    defn('1.4', '"Net Revenue"',
         'means Gross Revenue less: (a) refunds paid to End Customers; (b) chargebacks; and (c) Stripe '
         'processing fees and other payment processor fees directly attributable to collection.')
    defn('1.5', '"IP"',
         'means all intellectual property owned by Taurus relating to the Platform, including without '
         'limitation: source code, object code, software, algorithms, AI models, training data, prompt '
         'engineering, documentation, trademarks, trade names, service marks, logos, domain names, brand '
         'assets, trade secrets, and know-how.')
    defn('1.6', '"Territory"', 'means worldwide.')
    defn('1.7', '"Accounting Period"',
         'means each calendar month, unless otherwise agreed in writing by the Parties.')
    defn('1.8', '"Development Services"',
         'means software development, product management, AI research and development, quality assurance, '
         'infrastructure management, and other technical services provided by Taurus to maintain and '
         'improve the Platform.')

    # ── Article 2 — IP License ──
    art('ARTICLE 2 — IP LICENSE')
    sec('2.1', 'Grant of License.  Subject to the terms and conditions of this Agreement, Taurus hereby '
        'grants to VIGMIS US a non-exclusive, non-transferable, non-sublicensable license in the Territory to: '
        '(a) use the Platform solely for the purpose of marketing, demonstrating, and selling access to End '
        'Customers; (b) use Taurus\'s trademarks, brand names, and logos solely in connection with authorized '
        'marketing and sales activities; and (c) enter into agreements with End Customers for access to and '
        'use of the Platform.')
    sec('2.2', 'No Transfer of Ownership.  Nothing in this Agreement shall be construed to transfer any '
        'ownership interest in the IP from Taurus to VIGMIS US. All IP shall remain the exclusive property '
        'of Taurus. VIGMIS US shall acquire no rights in the IP except the limited license expressly granted herein.')
    sec('2.3', 'Sublicensing to End Customers.  VIGMIS US is authorized to grant End Customers a limited, '
        'non-exclusive right to access and use the Platform through its Terms of Service, consistent with '
        'the scope of this Agreement. All such rights granted to End Customers shall be subject to and '
        'limited by this Agreement.')
    sec('2.4', 'New Developments.  Any improvements, enhancements, or derivatives of the Platform or IP, '
        'whether developed by Taurus, VIGMIS US, or jointly, shall be owned exclusively by Taurus. VIGMIS US '
        'hereby assigns to Taurus all right, title, and interest in any such developments.')

    doc.add_page_break()

    # ── Article 3 — Services ──
    art('ARTICLE 3 — DEVELOPMENT AND INFRASTRUCTURE SERVICES')
    sec('3.1', 'Services Provided by Taurus.  During the Term, Taurus shall provide VIGMIS US with the '
        'following services necessary to operate the Platform: (a) Software Development — continuous '
        'development, enhancement, and maintenance of the Platform; (b) Infrastructure Management — hosting, '
        'cloud services (Railway, Supabase, Cloudflare, and equivalents), database management, and security; '
        '(c) AI Operations — management of AI providers (OpenAI, Anthropic, and others); (d) Product '
        'Management — feature prioritization, UX design, and product roadmap; (e) Technical Support — '
        'second-level escalated support; (f) Security — cybersecurity, data protection, and compliance.')
    sec('3.2', 'Service Standard.  Taurus shall provide the Services with reasonable skill and care, '
        'consistent with industry standards for SaaS platforms of similar nature and scale.')
    sec('3.3', 'Costs.  Each Party shall bear the costs allocated to it under the operating model adopted '
        'by the Parties from time to time.  As of the Effective Date: Taurus bears employee salaries, AI '
        'provider fees (OpenAI, Anthropic, and others), cloud infrastructure costs, R&D expenditures, and '
        'all other costs necessary to develop, maintain, and operate the Platform; VIGMIS US bears payment '
        'processor fees, digital advertising spend, CRM costs, and general administrative costs of operating '
        'VIGMIS US.  The Parties may adjust cost allocation by written amendment, provided any such '
        'adjustment maintains compliance with the arm\'s-length standard.')

    # ── Article 4 — Distribution ──
    art('ARTICLE 4 — DISTRIBUTION AND SALES RESPONSIBILITIES')
    sec('4.1', 'Sales and Marketing.  VIGMIS US shall be responsible for: (a) marketing the Platform to '
        'prospective End Customers in the Territory; (b) managing digital advertising channels (Google Ads, '
        'Meta Ads, TikTok Ads, and equivalents); (c) onboarding new End Customers; (d) managing customer '
        'relationships and CRM; (e) handling customer billing and payment collection through Stripe or '
        'equivalent payment processors; (f) managing refunds and chargebacks; (g) managing any affiliate '
        'or referral programs.')
    sec('4.2', 'Customer Contracts.  VIGMIS US shall enter into agreements with End Customers under its '
        'own name and on its own behalf, as the contracting entity. Such agreements shall be on terms no '
        'less protective of the Platform and IP than the protections afforded herein.')
    sec('4.3', 'Merchant of Record.  VIGMIS US shall be the merchant of record for all End Customer '
        'subscriptions and transactions processed through Stripe or any equivalent payment processor. '
        'VIGMIS US\'s name and identity shall appear on End Customer billing statements and payment '
        'receipts. VIGMIS US shall maintain the Stripe account in its own name and is responsible for '
        'all payment processing obligations, including chargebacks, refunds, and PCI compliance.')
    sec('4.4', 'VIGMIS US Costs.  VIGMIS US shall bear all costs related to its sales, marketing, and '
        'distribution activities, including payment processor fees, digital advertising spend, CRM costs, '
        'and general administrative costs of maintaining VIGMIS US as an operating entity.')

    # ── Article 5 — Revenue ──
    art('ARTICLE 5 — REVENUE SHARING AND PAYMENT')
    sec('5.1', 'Initial Revenue Allocation.  The Parties agree that the initial allocation of Net Revenue '
        'shall be as follows:')

    t_rev = doc.add_table(rows=1, cols=3)
    t_rev.style = 'Table Grid'
    t_rev.paragraph_format = None
    h_rev = t_rev.rows[0].cells
    for i, txt in enumerate(['Party', 'Percentage', 'Rationale']):
        h_rev[i].text = txt
        for run in h_rev[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for row_data in [
        ('Taurus', '75%', 'IP license royalty + Development Services fee'),
        ('VIGMIS US', '25%', 'Sales, marketing, distribution, and billing operations'),
    ]:
        add_table_row_en(t_rev, row_data)
    doc.add_paragraph()

    sec('5.2', 'Allocation Review.  The Parties may adjust the revenue allocation by written amendment, '
        'provided that: (a) any adjustment is supported by updated transfer pricing documentation; '
        '(b) the adjusted allocation remains consistent with the arm\'s-length standard under applicable '
        'Israeli and US tax law; and (c) any adjustment is reviewed by each Party\'s tax advisors prior '
        'to implementation.')
    sec('5.3', 'Arm\'s Length Justification.  The Parties acknowledge that the 75/25 allocation reflects '
        'an arm\'s-length arrangement: (a) Taurus contributes substantially all value-generating assets '
        '(IP, technology, development capability); (b) Taurus bears substantially all operating costs; '
        '(c) VIGMIS US\'s primary function is distribution and billing — a service that, on an arm\'s-length '
        'basis, would typically command a distribution margin of 20-30% of net revenues in the SaaS industry; '
        '(d) the 25% retained by VIGMIS US is designed to cover its actual costs plus a reasonable profit '
        'margin for its distribution function.')
    sec('5.4', 'Payment Mechanics.  Within fifteen (15) days following the end of each Accounting Period, '
        'VIGMIS US shall calculate Net Revenue and prepare a Revenue Statement. Within five (5) business '
        'days of preparing the Revenue Statement, VIGMIS US shall transfer 75% of Net Revenue to Taurus '
        'via wire transfer or ACH to Taurus\'s designated bank account. VIGMIS US shall retain 25%.')
    sec('5.5', 'Records and Audit.  VIGMIS US shall maintain accurate books and records relating to Gross '
        'Revenue, deductions, and Net Revenue. Taurus shall have the right, upon reasonable written notice, '
        'to audit such records no more than once per calendar year.')
    sec('5.6', 'Currency.  All amounts shall be calculated and paid in United States Dollars (USD). '
        'Where End Customer payments are collected in other currencies, conversion shall be at the exchange '
        'rate used by the applicable payment processor on the date of collection.')

    doc.add_page_break()

    # ── Article 6 — IP Ownership ──
    art('ARTICLE 6 — INTELLECTUAL PROPERTY OWNERSHIP AND PROTECTION')
    sec('6.1', 'Taurus IP.  As between the Parties, Taurus exclusively owns all IP, including without '
        'limitation: the VIGMIS source code, AI models, prompt libraries, customer acquisition methodologies, '
        'trademarks, domain names (vigmis.com and all related domains), logos, brand assets, and all '
        'proprietary workflows embedded in the Platform. VIGMIS US shall not take any action inconsistent '
        'with Taurus\'s ownership of the IP.')
    sec('6.2', 'New Developments.  Any improvements, enhancements, or derivatives of the Platform or IP, '
        'whether developed by Taurus, VIGMIS US, or jointly, shall be owned exclusively by Taurus. VIGMIS '
        'US hereby assigns to Taurus all right, title, and interest in any such developments.')
    sec('6.3', 'Brand Use.  VIGMIS US shall use the VIGMIS brand, trademarks, and logos only in the form '
        'and manner approved by Taurus, and only in connection with authorized activities under this Agreement.')
    sec('6.4', 'Notification.  VIGMIS US shall promptly notify Taurus of any actual or threatened '
        'infringement of the IP by any third party that comes to VIGMIS US\'s attention.')

    # ── Article 7 — Data Ownership ──
    art('ARTICLE 7 — DATA OWNERSHIP AND PRIVACY')
    sec('7.1', 'Data Ownership.  As between the Parties, ownership of data generated through the Platform '
        'is allocated as follows:')
    sub('a', 'Customer Account Data (business profiles, onboarding information, campaign settings): '
        'VIGMIS US holds the contractual relationship with End Customers; Taurus processes such data as a '
        'sub-processor to operate the Platform.')
    sub('b', 'Usage and Analytics Data (platform interactions, feature usage, performance metrics, '
        'aggregated statistics): owned by Taurus. Taurus may use such data to improve the Platform, '
        'train AI models (in anonymized form), and develop new features.')
    sub('c', 'AI Interaction Logs (prompts, AI-generated content, optimization decisions): owned by '
        'Taurus. May be used for model improvement and quality assurance in anonymized, aggregated form.')
    sub('d', 'Platform Infrastructure Data (logs, security events, system metrics): owned by Taurus.')
    sub('e', 'Billing and Payment Records: VIGMIS US retains records as merchant of record; Taurus '
        'retains copies as required for transfer pricing documentation and tax compliance.')
    sec('7.2', 'IP and Product Data.  For the avoidance of doubt, Taurus exclusively owns all IP as '
        'defined in Article 1.5, including the source code, AI models, prompt libraries, customer '
        'acquisition methodologies, and all proprietary workflows embedded in the Platform.')
    sec('7.3', 'Data Processing Agreement.  VIGMIS US shall ensure that its agreements with End Customers '
        'include appropriate data processing terms that permit Taurus to process End Customer data as a '
        'sub-processor, solely as necessary to operate and improve the Platform.')
    sec('7.4', 'Compliance.  Each Party shall comply with applicable data protection laws in its '
        'jurisdiction, including without limitation the GDPR (where applicable), Israel\'s Privacy '
        'Protection Law, and any applicable US federal or state privacy laws.')

    # ── Article 8 — Confidentiality ──
    art('ARTICLE 8 — CONFIDENTIALITY')
    sec('8.1', 'Confidential Information.  Each Party shall hold in strict confidence all Confidential '
        'Information of the other Party. "Confidential Information" means any non-public information '
        'disclosed by one Party to the other that is designated as confidential or that reasonably should '
        'be understood to be confidential, including financial information, customer data, technical '
        'information, and business plans.')
    sec('8.2', 'Exceptions.  The obligations of Section 8.1 do not apply to information that: (a) is or '
        'becomes publicly available through no breach of this Agreement; (b) was rightfully known by the '
        'receiving Party prior to disclosure; (c) is rightfully received from a third party without '
        'restriction; or (d) is required to be disclosed by law or court order, provided prompt prior '
        'notice is given.')
    sec('8.3', 'Survival.  Confidentiality obligations shall survive termination for five (5) years.')

    # ── Article 9 — Representations ──
    art('ARTICLE 9 — REPRESENTATIONS AND WARRANTIES')
    sec('9.1', 'Mutual.  Each Party represents and warrants that: (a) it has full legal authority to '
        'enter into this Agreement; (b) this Agreement constitutes a legal, valid, and binding obligation; '
        'and (c) its entry into and performance of this Agreement does not violate any applicable law or '
        'any agreement to which it is a party.')
    sec('9.2', 'Taurus.  Taurus represents and warrants that: (a) it owns the IP free and clear of any '
        'liens or encumbrances that would impair the license granted herein; (b) to its knowledge, the '
        'Platform does not infringe the intellectual property rights of any third party.')
    sec('9.3', 'VIGMIS US.  VIGMIS US represents and warrants that: (a) it shall conduct its sales and '
        'distribution activities in compliance with applicable law; (b) it shall not make representations '
        'about the Platform that are misleading or inconsistent with Taurus\'s documented specifications.')

    # ── Article 10 — Liability ──
    art('ARTICLE 10 — LIMITATION OF LIABILITY')
    sec('10.1', 'In no event shall either Party be liable to the other for any indirect, incidental, '
        'special, punitive, or consequential damages, even if advised of the possibility of such damages, '
        'arising out of or related to this Agreement.')
    sec('10.2', 'Each Party\'s aggregate liability to the other under or in connection with this Agreement '
        'shall not exceed the total Net Revenue paid or payable to Taurus in the twelve (12) months '
        'preceding the event giving rise to the claim.')

    doc.add_page_break()

    # ── Article 11 — Term and Termination ──
    art('ARTICLE 11 — TERM AND TERMINATION')
    sec('11.1', 'Term.  This Agreement shall commence on the Effective Date and shall continue indefinitely '
        'unless terminated as provided herein.')
    sec('11.2', 'Termination by Mutual Agreement.  The Parties may terminate this Agreement at any time '
        'by written agreement.')
    sec('11.3', 'Termination for Cause.  Either Party may terminate this Agreement upon thirty (30) days\' '
        'written notice if the other Party materially breaches this Agreement and fails to cure such breach '
        'within the notice period.')
    sec('11.4', 'Effect of Termination.  Upon termination: (a) the license granted under Article 2 shall '
        'immediately terminate; (b) VIGMIS US shall cease all use of the IP; (c) VIGMIS US shall promptly '
        'pay to Taurus all amounts due and outstanding; (d) each Party shall return or destroy the other\'s '
        'Confidential Information.')
    sec('11.5', 'Survival.  Articles 1, 6, 7, 8, 9, 10, 12, and 13 shall survive any termination or '
        'expiration of this Agreement.')
    sec('11.6', 'Insolvency and Wind-Down.  In the event of the dissolution, insolvency, bankruptcy, or '
        'winding-up of VIGMIS US (for any reason): (a) all existing End Customer relationships and contracts '
        'shall, to the extent permitted by applicable law, be assigned and transferred to Taurus or a '
        'Taurus-designated entity; (b) all End Customer data held or controlled by VIGMIS US shall be '
        'transferred to Taurus promptly and without charge; (c) the IP license granted under Article 2 '
        'shall immediately terminate; (d) VIGMIS US (or its liquidator, trustee, or equivalent) shall '
        'execute all documents reasonably necessary to effect the foregoing transfers; (e) any amounts '
        'owed to Taurus shall constitute a priority obligation in any VIGMIS US dissolution or insolvency '
        'proceeding, to the extent permitted by applicable law.')

    # ── Article 12 — Tax Cooperation ──
    art('ARTICLE 12 — TAX COOPERATION')
    sec('12.1', 'General Cooperation.  Each Party shall cooperate with the other Party\'s reasonable '
        'requests in connection with the preparation and filing of tax returns, tax audits, and regulatory '
        'inquiries related to the transactions contemplated by this Agreement.')
    sec('12.2', 'Form 5472 (US).  Taurus shall provide VIGMIS US with all information reasonably necessary '
        'for VIGMIS US to satisfy its US federal tax reporting obligations as a Foreign-Owned Single-Member '
        'LLC, including Form 5472 and the pro-forma Form 1120. VIGMIS US shall file all required forms '
        'timely. The Parties shall coordinate annual filings no later than thirty (30) days before the '
        'applicable deadline.')
    sec('12.3', 'CFC Reporting (Israel).  VIGMIS US shall provide Taurus with all information reasonably '
        'necessary for Taurus to assess and satisfy any Controlled Foreign Corporation (CFC) reporting '
        'obligations under Israeli tax law, including financial statements and revenue data for each fiscal year.')
    sec('12.4', 'Transfer Pricing Documentation.  The Parties shall maintain contemporaneous transfer '
        'pricing documentation sufficient to substantiate that the revenue allocation in Article 5 and '
        'cost allocation in Article 3.3 are consistent with the arm\'s-length standard. Such documentation '
        'shall be updated: (a) annually; (b) upon any material change in the operating model; and '
        '(c) upon any adjustment to the revenue allocation under Section 5.2.')
    sec('12.5', 'Document Retention.  Each Party shall retain all financial records, invoices, Revenue '
        'Statements, and supporting documentation relating to intercompany transactions for a minimum of '
        'seven (7) years, or such longer period as required by applicable law.')
    sec('12.6', 'No Tax Warranties.  Nothing in this Article shall be construed as either Party providing '
        'tax advice to the other. Each Party is solely responsible for obtaining independent tax advice '
        'and for its own tax compliance.')

    # ── Article 13 — General ──
    art('ARTICLE 13 — GENERAL PROVISIONS')
    sec('13.1', 'Governing Law.  This Agreement shall be governed by and construed in accordance with the '
        'laws of the State of Israel, without regard to its conflict of laws principles. The Parties agree '
        'to submit to the exclusive jurisdiction of the competent courts in Tel Aviv, Israel.')
    sec('13.2', 'Dispute Resolution.  The Parties shall attempt to resolve any dispute arising under this '
        'Agreement through good-faith negotiation. If the dispute cannot be resolved within thirty (30) '
        'days, the Parties may pursue legal remedies as permitted herein.')
    sec('13.3', 'Entire Agreement.  This Agreement constitutes the entire agreement between the Parties '
        'with respect to its subject matter and supersedes all prior and contemporaneous agreements, '
        'representations, and understandings.')
    sec('13.4', 'Amendment.  This Agreement may not be amended except by a written instrument signed by '
        'authorized representatives of both Parties.')
    sec('13.5', 'Waiver.  No waiver of any breach of this Agreement shall constitute a waiver of any '
        'subsequent breach.')
    sec('13.6', 'Severability.  If any provision of this Agreement is found to be invalid or unenforceable, '
        'the remaining provisions shall continue in full force and effect.')
    sec('13.7', 'Notices.  All notices shall be in writing and delivered to the addresses set forth in the '
        'preamble, by email with confirmation, courier, or registered mail.')
    sec('13.8', 'Relationship of Parties.  The Parties are independent contracting entities. Nothing in '
        'this Agreement shall be construed to create a partnership, joint venture, agency, employment, '
        'or fiduciary relationship between the Parties.')
    sec('13.9', 'Related Parties / Arm\'s Length.  The Parties acknowledge that they are related parties '
        'under common ownership. They confirm that the terms herein are intended to reflect an arm\'s-length '
        'arrangement designed to withstand scrutiny by Israeli and United States tax authorities.')
    sec('13.10', 'Counterparts.  This Agreement may be executed in counterparts, each of which shall be '
        'deemed an original, and all of which together shall constitute one and the same instrument. '
        'Electronic signatures shall be deemed valid.')

    add_hr_en(doc)

    # ── Signatures ──
    doc.add_heading('SIGNATURES', level=1)
    sig_note = doc.add_paragraph(
        'IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.'
    )
    sig_note.paragraph_format.space_after = Pt(20)

    sig_table = doc.add_table(rows=7, cols=2)
    sig_cells = [
        ['TAURUS MANAGEMENT AND INVESTMENTS LTD.', 'VIGMIS US LLC'],
        ['(Israeli private company, Reg. No. 514565118)', '(Wyoming LLC, wholly owned by Taurus)'],
        ['', ''],
        ['By: ___________________________', 'By: ___________________________'],
        ['Name: Amichai Krupik', 'Name: Amichai Krupik'],
        ['Title: Director / Authorized Signatory', 'Title: Manager / Authorized Signatory'],
        ['Date: ___________________________', 'Date: ___________________________'],
    ]
    for i, row_data in enumerate(sig_cells):
        row = sig_table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True
                if i == 1:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        'This document is a working draft prepared for internal business purposes and for review by '
        'qualified legal counsel and tax advisors in both Israel and the United States. It does not '
        'constitute legal or tax advice. The Parties should engage qualified counsel prior to '
        'finalizing and executing this Agreement.'
    )
    disclaimer.runs[0].italic = True
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(r'C:\vigmis\vigmis-main\LEGAL\INTERCOMPANY_AGREEMENT.docx')
    print('Agreement document saved')


# ── Run ───────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    build_hebrew_doc()
    build_agreement_doc()
    print('Done. Files saved to LEGAL folder.')
